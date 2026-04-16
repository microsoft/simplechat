# route_backend_conversation_export.py

import io
import json
import markdown2
import re
import tempfile
import zipfile
from collections import Counter, defaultdict
from datetime import datetime
from html import escape as _escape_html
from typing import Any, Dict, List, Optional

from bs4 import BeautifulSoup, NavigableString, Tag
from config import *
from flask import jsonify, make_response, request
from functions_appinsights import log_event
from functions_authentication import *
from functions_chat import sort_messages_by_thread
from functions_collaboration import (
    assert_user_can_view_collaboration_conversation,
    get_accessible_collaboration_message_thoughts,
    get_collaboration_conversation,
    is_collaboration_conversation,
    list_collaboration_messages,
)
from functions_conversation_metadata import update_conversation_with_metadata
from functions_debug import debug_print
from functions_message_artifacts import (
    build_message_artifact_payload_map,
    hydrate_agent_citations_from_artifacts,
    is_assistant_artifact_role,
)
from functions_settings import *
from functions_thoughts import get_thoughts_for_conversation
from swagger_wrapper import swagger_route, get_auth_security
from docx import Document as DocxDocument
from docx.shared import Inches, Pt


TRANSCRIPT_ROLES = {'user', 'assistant'}
SUMMARY_SOURCE_CHAR_LIMIT = 60000
DOCX_MARKDOWN_EXTRAS = ['fenced-code-blocks', 'tables', 'break-on-newline', 'cuddled-lists', 'strike']
EMAIL_SUBJECT_CHAR_LIMIT = 120
EMAIL_SUBJECT_SOURCE_CHAR_LIMIT = 12000


def register_route_backend_conversation_export(app):
    """Register conversation export API routes."""

    @app.route('/api/conversations/export', methods=['POST'])
    @swagger_route(security=get_auth_security())
    @login_required
    @user_required
    def api_export_conversations():
        """
        Export one or more conversations in JSON or Markdown format.
        Supports single-file or ZIP packaging.

        Request body:
            conversation_ids (list): List of conversation IDs to export.
            format (str): Export format — "json" or "markdown".
            packaging (str): Output packaging — "single" or "zip".
            include_summary_intro (bool): Whether to generate a per-conversation intro.
            summary_model_deployment (str): Optional model deployment for summary generation.
        """
        user_id = get_current_user_id()
        if not user_id:
            return jsonify({'error': 'User not authenticated'}), 401

        data = request.get_json(silent=True)
        if not data:
            return jsonify({'error': 'Request body is required'}), 400

        conversation_ids = data.get('conversation_ids', [])
        export_format = str(data.get('format', 'json')).lower()
        packaging = str(data.get('packaging', 'single')).lower()
        include_summary_intro = bool(data.get('include_summary_intro', False))
        summary_model_deployment = str(data.get('summary_model_deployment', '') or '').strip()

        if not conversation_ids or not isinstance(conversation_ids, list):
            return jsonify({'error': 'At least one conversation_id is required'}), 400

        if export_format not in ('json', 'markdown', 'pdf'):
            return jsonify({'error': 'Format must be "json", "markdown", or "pdf"'}), 400

        if packaging not in ('single', 'zip'):
            return jsonify({'error': 'Packaging must be "single" or "zip"'}), 400

        try:
            settings = get_settings()
            exported = []
            for conv_id in conversation_ids:
                conversation = None
                messages = []
                try:
                    conversation = cosmos_conversations_container.read_item(
                        item=conv_id,
                        partition_key=conv_id
                    )
                    if conversation.get('user_id') != user_id:
                        debug_print(f"Export: user {user_id} does not own conversation {conv_id}")
                        continue

                    message_query = """
                        SELECT * FROM c
                        WHERE c.conversation_id = @conversation_id
                        ORDER BY c.timestamp ASC
                    """
                    messages = list(cosmos_messages_container.query_items(
                        query=message_query,
                        parameters=[{'name': '@conversation_id', 'value': conv_id}],
                        partition_key=conv_id
                    ))
                except Exception:
                    try:
                        conversation = get_collaboration_conversation(conv_id)
                        access_context = assert_user_can_view_collaboration_conversation(
                            user_id,
                            conversation,
                            allow_pending=True,
                        )
                        user_state = access_context.get('user_state') or {}
                        conversation = dict(conversation)
                        conversation['is_pinned'] = bool(user_state.get('is_pinned', False))
                        conversation['is_hidden'] = bool(user_state.get('is_hidden', False))
                        messages = list_collaboration_messages(conv_id)
                    except Exception:
                        debug_print(f"Export: conversation {conv_id} not found or access denied")
                        continue

                exported.append(
                    _build_export_entry(
                        conversation=conversation,
                        raw_messages=messages,
                        user_id=user_id,
                        settings=settings,
                        include_summary_intro=include_summary_intro,
                        summary_model_deployment=summary_model_deployment
                    )
                )

            if not exported:
                return jsonify({'error': 'No accessible conversations found'}), 404

            timestamp_str = datetime.utcnow().strftime('%Y%m%d_%H%M%S')

            if packaging == 'zip':
                return _build_zip_response(exported, export_format, timestamp_str)

            return _build_single_file_response(exported, export_format, timestamp_str)

        except Exception as exc:
            debug_print(f"Export error: {str(exc)}")
            log_event(f"Conversation export failed: {exc}", level="WARNING")
            return jsonify({'error': f'Export failed: {str(exc)}'}), 500

    @app.route('/api/message/export-word', methods=['POST'])
    @swagger_route(security=get_auth_security())
    @login_required
    @user_required
    def api_export_message_word():
        """
        Export a single message as a Word (.docx) document.

        Request body:
            message_id (str): ID of the message to export.
            conversation_id (str): ID of the conversation the message belongs to.
        """
        user_id = get_current_user_id()
        if not user_id:
            return jsonify({'error': 'User not authenticated'}), 401

        data = request.get_json(silent=True)
        if not data:
            return jsonify({'error': 'Request body is required'}), 400

        message_id = str(data.get('message_id', '') or '').strip()
        conversation_id = str(data.get('conversation_id', '') or '').strip()

        if not message_id or not conversation_id:
            return jsonify({'error': 'message_id and conversation_id are required'}), 400

        try:
            message = _load_export_message_for_user(
                user_id=user_id,
                conversation_id=conversation_id,
                message_id=message_id
            )

            document_bytes = _message_to_docx_bytes(message)
            timestamp_str = datetime.utcnow().strftime('%Y%m%d_%H%M%S')
            filename = f"message_export_{timestamp_str}.docx"

            response = make_response(document_bytes)
            response.headers['Content-Type'] = (
                'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
            response.headers['Content-Disposition'] = f'attachment; filename="{filename}"'
            return response

        except LookupError as exc:
            return jsonify({'error': str(exc)}), 404
        except PermissionError as exc:
            return jsonify({'error': str(exc)}), 403

        except Exception as exc:
            debug_print(f"Message export error: {str(exc)}")
            log_event(f"Message export failed: {exc}", level="WARNING")
            return jsonify({'error': 'Export failed due to a server error. Please try again later.'}), 500

    @app.route('/api/message/export-email-draft', methods=['POST'])
    @swagger_route(security=get_auth_security())
    @login_required
    @user_required
    def api_export_message_email_draft():
        """
        Build a mailto-ready email draft for a single message.

        Request body:
            message_id (str): ID of the message to export.
            conversation_id (str): ID of the conversation the message belongs to.
            summary_model_deployment (str): Optional model deployment for subject generation.
        """
        user_id = get_current_user_id()
        if not user_id:
            return jsonify({'error': 'User not authenticated'}), 401

        data = request.get_json(silent=True)
        if not data:
            return jsonify({'error': 'Request body is required'}), 400

        message_id = str(data.get('message_id', '') or '').strip()
        conversation_id = str(data.get('conversation_id', '') or '').strip()
        summary_model_deployment = str(data.get('summary_model_deployment', '') or '').strip()

        if not message_id or not conversation_id:
            return jsonify({'error': 'message_id and conversation_id are required'}), 400

        try:
            settings = get_settings()
            message = _load_export_message_for_user(
                user_id=user_id,
                conversation_id=conversation_id,
                message_id=message_id
            )
            draft_payload = _message_to_email_draft_payload(
                message=message,
                settings=settings,
                summary_model_deployment=summary_model_deployment
            )
            return jsonify(draft_payload), 200

        except LookupError as exc:
            return jsonify({'error': str(exc)}), 404
        except PermissionError as exc:
            return jsonify({'error': str(exc)}), 403

        except Exception as exc:
            debug_print(f"Message email draft export error: {str(exc)}")
            log_event(f"Message email draft export failed: {exc}", level="WARNING")
            return jsonify({'error': 'Email draft export failed due to a server error. Please try again later.'}), 500


def _build_export_entry(
    conversation: Dict[str, Any],
    raw_messages: List[Dict[str, Any]],
    user_id: str,
    settings: Dict[str, Any],
    include_summary_intro: bool = False,
    summary_model_deployment: str = ''
) -> Dict[str, Any]:
    artifact_payload_map = build_message_artifact_payload_map(raw_messages)
    filtered_messages = _filter_messages_for_export(raw_messages)
    filtered_messages = hydrate_agent_citations_from_artifacts(filtered_messages, artifact_payload_map)
    ordered_messages = sort_messages_by_thread(filtered_messages)

    raw_thoughts = [] if is_collaboration_conversation(conversation) else get_thoughts_for_conversation(conversation.get('id'), user_id)
    thoughts_by_message = defaultdict(list)
    for thought in raw_thoughts:
        thoughts_by_message[thought.get('message_id')].append(_sanitize_thought(thought))

    exported_messages = []
    role_counts = Counter()
    total_citation_counts = Counter({'document': 0, 'web': 0, 'agent_tool': 0, 'legacy': 0, 'total': 0})
    transcript_index = 0
    total_thoughts = 0

    for sequence_index, message in enumerate(ordered_messages, start=1):
        role = message.get('role', 'unknown')
        role_counts[role] += 1

        message_transcript_index = None
        if role in TRANSCRIPT_ROLES:
            transcript_index += 1
            message_transcript_index = transcript_index

        thoughts = thoughts_by_message.get(message.get('id'), [])
        if not thoughts and is_collaboration_conversation(conversation):
            collaboration_thoughts = get_accessible_collaboration_message_thoughts(
                conversation,
                message,
                user_id,
            )
            thoughts = [_sanitize_thought(thought) for thought in collaboration_thoughts]
        exported_message = _sanitize_message(
            message,
            sequence_index=sequence_index,
            transcript_index=message_transcript_index,
            thoughts=thoughts
        )
        exported_messages.append(exported_message)

        counts = exported_message.get('citation_counts', {})
        for key in total_citation_counts:
            total_citation_counts[key] += counts.get(key, 0)
        total_thoughts += len(thoughts)

    # Compute message time range for summary caching
    message_time_start = None
    message_time_end = None
    if ordered_messages:
        message_time_start = ordered_messages[0].get('timestamp')
        message_time_end = ordered_messages[-1].get('timestamp')

    sanitized_conversation = _sanitize_conversation(
        conversation,
        messages=exported_messages,
        role_counts=role_counts,
        citation_counts=total_citation_counts,
        thought_count=total_thoughts
    )
    summary_intro = _build_summary_intro(
        messages=exported_messages,
        conversation=conversation,
        sanitized_conversation=sanitized_conversation,
        settings=settings,
        enabled=include_summary_intro,
        summary_model_deployment=summary_model_deployment,
        message_time_start=message_time_start,
        message_time_end=message_time_end
    )

    return {
        'conversation': sanitized_conversation,
        'summary_intro': summary_intro,
        'messages': exported_messages
    }


def _filter_messages_for_export(messages: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    filtered_messages = []
    for message in messages:
        if is_assistant_artifact_role(message.get('role')):
            continue

        metadata = message.get('metadata', {}) or {}
        if metadata.get('is_deleted') is True:
            continue

        thread_info = metadata.get('thread_info', {}) or {}
        active = thread_info.get('active_thread')
        if active is True or active is None or 'active_thread' not in thread_info:
            filtered_messages.append(message)

    return filtered_messages


def _sanitize_conversation(
    conversation: Dict[str, Any],
    messages: List[Dict[str, Any]],
    role_counts: Counter,
    citation_counts: Counter,
    thought_count: int
) -> Dict[str, Any]:
    transcript_count = sum(1 for message in messages if message.get('is_transcript_message'))
    return {
        'id': conversation.get('id'),
        'title': conversation.get('title', 'Untitled'),
        'last_updated': conversation.get('last_updated') or conversation.get('updated_at', ''),
        'chat_type': conversation.get('chat_type', 'personal'),
        'tags': conversation.get('tags', []),
        'context': conversation.get('context', []),
        'classification': conversation.get('classification', []),
        'strict': conversation.get('strict', False),
        'is_pinned': conversation.get('is_pinned', False),
        'scope_locked': conversation.get('scope_locked'),
        'locked_contexts': conversation.get('locked_contexts', []),
        'message_count': len(messages),
        'transcript_message_count': transcript_count,
        'message_counts_by_role': dict(role_counts),
        'citation_counts': dict(citation_counts),
        'thought_count': thought_count
    }


def _sanitize_message(
    message: Dict[str, Any],
    sequence_index: int,
    transcript_index: Optional[int],
    thoughts: List[Dict[str, Any]]
) -> Dict[str, Any]:
    role = message.get('role', '')
    content = message.get('content', '')
    raw_citation_buckets = _collect_raw_citation_buckets(message)
    normalized_citations = _normalize_citations(raw_citation_buckets)
    citation_counts = _build_citation_counts(normalized_citations)
    details = _curate_message_details(message, citation_counts, len(thoughts))

    return {
        'id': message.get('id'),
        'role': role,
        'speaker_label': _role_to_label(role),
        'sequence_index': sequence_index,
        'transcript_index': transcript_index,
        'label': f"Turn {transcript_index}" if transcript_index else f"Message {sequence_index}",
        'is_transcript_message': role in TRANSCRIPT_ROLES,
        'timestamp': message.get('timestamp', ''),
        'content': content,
        'content_text': _normalize_content(content),
        'details': details,
        'citations': normalized_citations,
        'citation_counts': citation_counts,
        'thoughts': thoughts,
        'legacy_citations': raw_citation_buckets['legacy'],
        'hybrid_citations': raw_citation_buckets['hybrid'],
        'web_search_citations': raw_citation_buckets['web'],
        'agent_citations': raw_citation_buckets['agent']
    }


def _sanitize_thought(thought: Dict[str, Any]) -> Dict[str, Any]:
    return {
        'step_index': thought.get('step_index'),
        'step_type': thought.get('step_type'),
        'content': thought.get('content'),
        'detail': thought.get('detail'),
        'duration_ms': thought.get('duration_ms'),
        'timestamp': thought.get('timestamp')
    }


def _collect_raw_citation_buckets(message: Dict[str, Any]) -> Dict[str, List[Any]]:
    def ensure_list(value: Any) -> List[Any]:
        if not value:
            return []
        return value if isinstance(value, list) else [value]

    return {
        'legacy': ensure_list(message.get('citations')),
        'hybrid': ensure_list(message.get('hybrid_citations')),
        'web': ensure_list(message.get('web_search_citations')),
        'agent': ensure_list(message.get('agent_citations'))
    }


def _normalize_citations(raw_citation_buckets: Dict[str, List[Any]]) -> List[Dict[str, Any]]:
    normalized = []

    for citation in raw_citation_buckets.get('hybrid', []):
        if isinstance(citation, dict):
            normalized.append({
                'citation_type': 'document',
                'label': _build_document_citation_label(citation),
                'file_name': citation.get('file_name'),
                'title': citation.get('title') or citation.get('file_name'),
                'page_number': citation.get('page_number'),
                'citation_id': citation.get('citation_id'),
                'chunk_id': citation.get('chunk_id'),
                'metadata_type': citation.get('metadata_type'),
                'metadata_content': citation.get('metadata_content'),
                'score': citation.get('score'),
                'classification': citation.get('classification'),
                'url': citation.get('url')
            })
        else:
            normalized.append({
                'citation_type': 'document',
                'label': str(citation),
                'value': str(citation)
            })

    for citation in raw_citation_buckets.get('web', []):
        if isinstance(citation, dict):
            title = citation.get('title') or citation.get('url') or 'Web source'
            normalized.append({
                'citation_type': 'web',
                'label': title,
                'title': title,
                'url': citation.get('url')
            })
        else:
            normalized.append({
                'citation_type': 'web',
                'label': str(citation),
                'value': str(citation)
            })

    for citation in raw_citation_buckets.get('agent', []):
        if isinstance(citation, dict):
            tool_name = citation.get('tool_name') or citation.get('function_name') or 'Tool invocation'
            normalized.append({
                'citation_type': 'agent_tool',
                'label': tool_name,
                'tool_name': citation.get('tool_name'),
                'function_name': citation.get('function_name'),
                'plugin_name': citation.get('plugin_name'),
                'success': citation.get('success'),
                'timestamp': citation.get('timestamp')
            })
        else:
            normalized.append({
                'citation_type': 'agent_tool',
                'label': str(citation),
                'value': str(citation)
            })

    for citation in raw_citation_buckets.get('legacy', []):
        if isinstance(citation, dict):
            title = citation.get('title') or citation.get('filepath') or citation.get('url') or 'Legacy citation'
            normalized.append({
                'citation_type': 'legacy',
                'label': title,
                'title': title,
                'url': citation.get('url'),
                'filepath': citation.get('filepath')
            })
        else:
            normalized.append({
                'citation_type': 'legacy',
                'label': str(citation),
                'value': str(citation)
            })

    return normalized


def _build_document_citation_label(citation: Dict[str, Any]) -> str:
    file_name = citation.get('file_name') or citation.get('title') or 'Document source'
    metadata_type = citation.get('metadata_type')
    page_number = citation.get('page_number')

    if metadata_type:
        return f"{file_name} — {metadata_type.replace('_', ' ').title()}"
    if page_number not in (None, ''):
        return f"{file_name} — Page {page_number}"
    return file_name


def _build_citation_counts(citations: List[Dict[str, Any]]) -> Dict[str, int]:
    counts = {
        'document': 0,
        'web': 0,
        'agent_tool': 0,
        'legacy': 0,
        'total': len(citations)
    }
    for citation in citations:
        citation_type = citation.get('citation_type')
        if citation_type in counts:
            counts[citation_type] += 1
    return counts


def _curate_message_details(
    message: Dict[str, Any],
    citation_counts: Dict[str, int],
    thought_count: int
) -> Dict[str, Any]:
    role = message.get('role', '')
    metadata = message.get('metadata', {}) or {}
    details: Dict[str, Any] = {}

    if role == 'user':
        details['interaction_mode'] = _remove_empty_values({
            'button_states': metadata.get('button_states'),
            'workspace_search': _curate_workspace_search(metadata.get('workspace_search')),
            'prompt_selection': _curate_prompt_selection(metadata.get('prompt_selection')),
            'agent_selection': _curate_agent_selection(metadata.get('agent_selection')),
            'model_selection': _curate_model_selection(metadata.get('model_selection'))
        })
    elif role == 'assistant':
        details['generation'] = _remove_empty_values({
            'augmented': message.get('augmented'),
            'model_deployment': message.get('model_deployment_name'),
            'agent_name': message.get('agent_name'),
            'agent_display_name': message.get('agent_display_name'),
            'reasoning_effort': metadata.get('reasoning_effort'),
            'hybrid_search_query': message.get('hybridsearch_query'),
            'token_usage': _curate_token_usage(metadata.get('token_usage')),
            'citation_counts': citation_counts,
            'thought_count': thought_count
        })
    else:
        details['message_context'] = _remove_empty_values({
            'filename': message.get('filename'),
            'prompt': message.get('prompt'),
            'is_table': message.get('is_table'),
            'model_deployment': message.get('model_deployment_name')
        })

    return _remove_empty_values(details)


def _curate_workspace_search(workspace_search: Optional[Dict[str, Any]]) -> Dict[str, Any]:
    if not isinstance(workspace_search, dict):
        return {}
    return _remove_empty_values({
        'search_enabled': workspace_search.get('search_enabled'),
        'document_scope': workspace_search.get('document_scope'),
        'document_name': workspace_search.get('document_name'),
        'document_filename': workspace_search.get('document_filename'),
        'group_name': workspace_search.get('group_name'),
        'classification': workspace_search.get('classification'),
        'public_workspace_id': workspace_search.get('active_public_workspace_id')
    })


def _curate_prompt_selection(prompt_selection: Optional[Dict[str, Any]]) -> Dict[str, Any]:
    if not isinstance(prompt_selection, dict):
        return {}
    return _remove_empty_values({
        'prompt_name': prompt_selection.get('prompt_name'),
        'selected_prompt_index': prompt_selection.get('selected_prompt_index'),
        'selected_prompt_text': prompt_selection.get('selected_prompt_text')
    })


def _curate_agent_selection(agent_selection: Optional[Dict[str, Any]]) -> Dict[str, Any]:
    if not isinstance(agent_selection, dict):
        return {}
    return _remove_empty_values({
        'selected_agent': agent_selection.get('selected_agent'),
        'agent_display_name': agent_selection.get('agent_display_name'),
        'is_global': agent_selection.get('is_global'),
        'is_group': agent_selection.get('is_group'),
        'group_name': agent_selection.get('group_name')
    })


def _curate_model_selection(model_selection: Optional[Dict[str, Any]]) -> Dict[str, Any]:
    if not isinstance(model_selection, dict):
        return {}
    return _remove_empty_values({
        'selected_model': model_selection.get('selected_model'),
        'frontend_requested_model': model_selection.get('frontend_requested_model'),
        'reasoning_effort': model_selection.get('reasoning_effort'),
        'streaming': model_selection.get('streaming')
    })


def _curate_token_usage(token_usage: Any) -> Dict[str, Any]:
    if not isinstance(token_usage, dict):
        return {}
    return _remove_empty_values({
        'prompt_tokens': token_usage.get('prompt_tokens'),
        'completion_tokens': token_usage.get('completion_tokens'),
        'total_tokens': token_usage.get('total_tokens')
    })


def _remove_empty_values(value: Any) -> Any:
    if isinstance(value, dict):
        cleaned = {}
        for key, item in value.items():
            cleaned_item = _remove_empty_values(item)
            if cleaned_item in (None, '', [], {}):
                continue
            cleaned[key] = cleaned_item
        return cleaned

    if isinstance(value, list):
        cleaned_list = []
        for item in value:
            cleaned_item = _remove_empty_values(item)
            if cleaned_item in (None, '', [], {}):
                continue
            cleaned_list.append(cleaned_item)
        return cleaned_list

    return value


def generate_conversation_summary(
    messages: List[Dict[str, Any]],
    conversation_title: str,
    settings: Dict[str, Any],
    model_deployment: str,
    message_time_start: str = None,
    message_time_end: str = None,
    conversation_id: str = None
) -> Dict[str, Any]:
    """Generate a conversation summary using the LLM and optionally persist it.

    This is the shared helper used by both the export pipeline and the
    on-demand summary API endpoint.  Returns a summary dict suitable for
    storage in conversation metadata.

    Raises ValueError when there is no content to summarise and
    RuntimeError on model errors.
    """
    transcript_lines = []
    for message in messages:
        content_text = message.get('content_text', '')
        if not content_text:
            continue
        role = message.get('role', 'unknown')
        speaker = message.get('speaker_label', role).upper()
        transcript_lines.append(f"{speaker}: {content_text}")

    transcript_text = '\n\n'.join(transcript_lines).strip()
    if not transcript_text:
        raise ValueError('No message content was available to summarize.')

    transcript_text = _truncate_for_summary(transcript_text)

    gpt_client, gpt_model = _initialize_gpt_client(settings, model_deployment)
    summary_prompt = (
        "You are summarizing a conversation for an export document. "
        "Read the full conversation below and write a concise summary. "
        "Use your judgement on length: for short conversations write one brief paragraph, "
        "for longer or more detailed conversations write two paragraphs. "
        "If you need refer to the user, use their name, but do not refer to the user too often."
        "Cover the goals, the key topics discussed, any data or tools referenced, "
        "and the main outcomes or answers provided. "
        "Be factual and neutral. Return plain text only — no headings, no bullet points, no markdown formatting."
    )

    model_lower = gpt_model.lower()
    is_reasoning_model = (
        'o1' in model_lower or 'o3' in model_lower or 'gpt-5' in model_lower
    )
    instruction_role = 'developer' if is_reasoning_model else 'system'

    debug_print(f"Summary generation: sending {len(transcript_lines)} messages "
                f"({len(transcript_text)} chars) to {gpt_model} (role={instruction_role})")

    summary_response = gpt_client.chat.completions.create(
        model=gpt_model,
        messages=[
            {
                'role': instruction_role,
                'content': summary_prompt
            },
            {
                'role': 'user',
                'content': (
                    f"Conversation Title: {conversation_title}\n\n"
                    f"{transcript_text}"
                )
            }
        ]
    )

    debug_print(f"Summary generation: response choices="
                f"{len(summary_response.choices) if summary_response.choices else 0}, "
                f"finish_reason={summary_response.choices[0].finish_reason if summary_response.choices else 'N/A'}")

    summary_text = (summary_response.choices[0].message.content or '').strip() if summary_response.choices else ''
    if not summary_text:
        debug_print('Summary generation: model returned an empty response')
        log_event('Conversation summary generation returned empty response', level='WARNING')
        raise RuntimeError('Summary model returned an empty response.')

    summary_data = {
        'content': summary_text,
        'model_deployment': gpt_model,
        'generated_at': datetime.utcnow().isoformat(),
        'message_time_start': message_time_start,
        'message_time_end': message_time_end
    }

    # Persist to Cosmos when a conversation_id is available
    if conversation_id:
        try:
            update_conversation_with_metadata(conversation_id, {'summary': summary_data})
            debug_print(f"Summary persisted to conversation {conversation_id}")
        except Exception as persist_exc:
            debug_print(f"Failed to persist summary to Cosmos: {persist_exc}")
            log_event(f"Failed to persist conversation summary: {persist_exc}", level="WARNING")

    return summary_data


def _build_summary_intro(
    messages: List[Dict[str, Any]],
    conversation: Dict[str, Any],
    sanitized_conversation: Dict[str, Any],
    settings: Dict[str, Any],
    enabled: bool,
    summary_model_deployment: str,
    message_time_start: str = None,
    message_time_end: str = None
) -> Dict[str, Any]:
    """Build the summary_intro block for the export payload.

    Uses cached summary from conversation metadata when present and
    still current (no newer messages).  Otherwise generates a fresh
    summary via ``generate_conversation_summary`` and persists it.
    """
    summary_intro = {
        'enabled': enabled,
        'generated': False,
        'model_deployment': summary_model_deployment or None,
        'generated_at': None,
        'content': '',
        'error': None
    }

    if not enabled:
        return summary_intro

    # Check for a cached summary stored in the conversation document
    existing_summary = conversation.get('summary')
    if existing_summary and isinstance(existing_summary, dict):
        cached_end = existing_summary.get('message_time_end')
        if cached_end and message_time_end and cached_end >= message_time_end:
            debug_print('Export summary: using cached summary from conversation metadata')
            summary_intro.update({
                'generated': True,
                'model_deployment': existing_summary.get('model_deployment'),
                'generated_at': existing_summary.get('generated_at'),
                'content': existing_summary.get('content', ''),
                'error': None
            })
            return summary_intro
        debug_print('Export summary: cached summary is stale, regenerating')

    try:
        conversation_id = conversation.get('id')
        conversation_title = sanitized_conversation.get('title', 'Untitled')

        summary_data = generate_conversation_summary(
            messages=messages,
            conversation_title=conversation_title,
            settings=settings,
            model_deployment=summary_model_deployment,
            message_time_start=message_time_start,
            message_time_end=message_time_end,
            conversation_id=conversation_id
        )

        summary_intro.update({
            'generated': True,
            'model_deployment': summary_data.get('model_deployment'),
            'generated_at': summary_data.get('generated_at'),
            'content': summary_data.get('content', ''),
            'error': None
        })
        return summary_intro

    except (ValueError, RuntimeError) as known_exc:
        debug_print(f"Export summary generation issue: {known_exc}")
        summary_intro['error'] = str(known_exc)
        if hasattr(known_exc, 'model_deployment'):
            summary_intro['model_deployment'] = known_exc.model_deployment
        return summary_intro

    except Exception as exc:
        debug_print(f"Export summary generation failed: {exc}")
        log_event(f"Conversation export summary generation failed: {exc}", level="WARNING")
        summary_intro['error'] = str(exc)
        return summary_intro


def _truncate_for_summary(transcript_text: str) -> str:
    if len(transcript_text) <= SUMMARY_SOURCE_CHAR_LIMIT:
        return transcript_text

    head_chars = SUMMARY_SOURCE_CHAR_LIMIT // 2
    tail_chars = SUMMARY_SOURCE_CHAR_LIMIT - head_chars
    return (
        transcript_text[:head_chars]
        + "\n\n[... transcript truncated for export summary generation ...]\n\n"
        + transcript_text[-tail_chars:]
    )


def _initialize_gpt_client(settings: Dict[str, Any], requested_model: str = ''):
    enable_gpt_apim = settings.get('enable_gpt_apim', False)

    if enable_gpt_apim:
        raw_models = settings.get('azure_apim_gpt_deployment', '') or ''
        apim_models = [model.strip() for model in raw_models.split(',') if model.strip()]
        if not apim_models:
            raise ValueError('APIM GPT deployment name is not configured.')

        if requested_model and requested_model not in apim_models:
            raise ValueError(f"Requested summary model '{requested_model}' is not configured for APIM.")

        gpt_model = requested_model or apim_models[0]
        gpt_client = AzureOpenAI(
            api_version=settings.get('azure_apim_gpt_api_version'),
            azure_endpoint=settings.get('azure_apim_gpt_endpoint'),
            api_key=settings.get('azure_apim_gpt_subscription_key')
        )
        return gpt_client, gpt_model

    auth_type = settings.get('azure_openai_gpt_authentication_type')
    endpoint = settings.get('azure_openai_gpt_endpoint')
    api_version = settings.get('azure_openai_gpt_api_version')
    gpt_model_obj = settings.get('gpt_model', {}) or {}

    if requested_model:
        gpt_model = requested_model
    elif gpt_model_obj.get('selected'):
        gpt_model = gpt_model_obj['selected'][0]['deploymentName']
    else:
        raise ValueError('No GPT model selected or configured for export summary generation.')

    if auth_type == 'managed_identity':
        token_provider = get_bearer_token_provider(DefaultAzureCredential(), cognitive_services_scope)
        gpt_client = AzureOpenAI(
            api_version=api_version,
            azure_endpoint=endpoint,
            azure_ad_token_provider=token_provider
        )
    else:
        api_key = settings.get('azure_openai_gpt_key')
        if not api_key:
            raise ValueError('Azure OpenAI API Key not configured.')
        gpt_client = AzureOpenAI(
            api_version=api_version,
            azure_endpoint=endpoint,
            api_key=api_key
        )

    return gpt_client, gpt_model


def _build_single_file_response(exported: List[Dict[str, Any]], export_format: str, timestamp_str: str):
    """Build a single-file download response."""
    if export_format == 'json':
        content = json.dumps(exported, indent=2, ensure_ascii=False, default=str)
        filename = f"conversations_export_{timestamp_str}.json"
        content_type = 'application/json; charset=utf-8'
    elif export_format == 'pdf':
        if len(exported) == 1:
            content = _conversation_to_pdf_bytes(exported[0])
        else:
            combined_parts = []
            for idx, entry in enumerate(exported):
                if idx > 0:
                    combined_parts.append(
                        '<div style="margin-top: 24pt; border-top: 2px solid #999; '
                        'padding-top: 12pt;"></div>'
                    )
                combined_parts.append(_build_pdf_html_body(entry))
            content = _html_body_to_pdf_bytes('\n'.join(combined_parts))
        filename = f"conversations_export_{timestamp_str}.pdf"
        content_type = 'application/pdf'
    else:
        parts = []
        for entry in exported:
            parts.append(_conversation_to_markdown(entry))
        content = '\n\n---\n\n'.join(parts)
        filename = f"conversations_export_{timestamp_str}.md"
        content_type = 'text/markdown; charset=utf-8'

    response = make_response(content)
    response.headers['Content-Type'] = content_type
    response.headers['Content-Disposition'] = f'attachment; filename="{filename}"'
    return response


def _build_zip_response(exported: List[Dict[str, Any]], export_format: str, timestamp_str: str):
    """Build a ZIP archive containing one file per conversation."""
    buffer = io.BytesIO()
    with zipfile.ZipFile(buffer, 'w', zipfile.ZIP_DEFLATED) as zf:
        for entry in exported:
            conversation = entry['conversation']
            safe_title = _safe_filename(conversation.get('title', 'Untitled'))
            conversation_id_short = conversation.get('id', 'unknown')[:8]

            if export_format == 'json':
                file_content = json.dumps(entry, indent=2, ensure_ascii=False, default=str)
                ext = 'json'
            elif export_format == 'pdf':
                file_content = _conversation_to_pdf_bytes(entry)
                ext = 'pdf'
            else:
                file_content = _conversation_to_markdown(entry)
                ext = 'md'

            file_name = f"{safe_title}_{conversation_id_short}.{ext}"
            zf.writestr(file_name, file_content)

    buffer.seek(0)
    filename = f"conversations_export_{timestamp_str}.zip"

    response = make_response(buffer.read())
    response.headers['Content-Type'] = 'application/zip'
    response.headers['Content-Disposition'] = f'attachment; filename="{filename}"'
    return response


def _conversation_to_markdown(entry: Dict[str, Any]) -> str:
    """Convert a conversation + messages entry to Markdown format."""
    conversation = entry['conversation']
    messages = entry['messages']
    summary_intro = entry.get('summary_intro', {}) or {}

    transcript_messages = [message for message in messages if message.get('is_transcript_message')]
    detail_messages = [message for message in messages if message.get('details')]
    reference_messages = [message for message in messages if message.get('citations')]
    thought_messages = [message for message in messages if message.get('thoughts')]
    supplemental_messages = [message for message in messages if not message.get('is_transcript_message')]

    lines: List[str] = []
    lines.append(f"# {conversation.get('title', 'Untitled')}")
    lines.append('')
    lines.append(f"**Last Updated:** {conversation.get('last_updated', '')}  ")
    lines.append(f"**Chat Type:** {conversation.get('chat_type', 'personal')}  ")
    lines.append(f"**Messages:** {conversation.get('message_count', len(messages))}  ")
    if conversation.get('tags'):
        lines.append(f"**Tags:** {', '.join(_format_tag(tag) for tag in conversation.get('tags', []))}  ")
    if conversation.get('classification'):
        lines.append(f"**Classification:** {', '.join(_format_tag(item) for item in conversation.get('classification', []))}  ")
    lines.append('')

    if summary_intro.get('enabled') and summary_intro.get('generated') and summary_intro.get('content'):
        lines.append('## Abstract')
        lines.append('')
        lines.append(summary_intro.get('content', ''))
        lines.append('')
        lines.append(f"_Generated with {summary_intro.get('model_deployment') or 'configured model'} on {summary_intro.get('generated_at')}_")
        lines.append('')
    elif summary_intro.get('enabled') and summary_intro.get('error'):
        lines.append('> _A summary intro was requested, but it could not be generated for this export._')
        lines.append(f"> _Error: {summary_intro.get('error')}_")
        lines.append('')

    lines.append('## Transcript')
    lines.append('')
    if not transcript_messages:
        lines.append('_No user or assistant transcript messages were available for export._')
        lines.append('')
    else:
        for message in transcript_messages:
            lines.append(f"### {message.get('label')} — {message.get('speaker_label')}")
            if message.get('timestamp'):
                lines.append(f"*{message.get('timestamp')}*")
            lines.append('')
            lines.append(message.get('content_text') or '_No content recorded._')
            lines.append('')

    lines.append('## Appendix A — Conversation Metadata')
    lines.append('')
    metadata_to_render = _remove_empty_values({
        'context': conversation.get('context'),
        'classification': conversation.get('classification'),
        'strict': conversation.get('strict'),
        'is_pinned': conversation.get('is_pinned'),
        'scope_locked': conversation.get('scope_locked'),
        'locked_contexts': conversation.get('locked_contexts'),
        'message_counts_by_role': conversation.get('message_counts_by_role'),
        'citation_counts': conversation.get('citation_counts'),
        'thought_count': conversation.get('thought_count')
    })
    _append_markdown_mapping(lines, metadata_to_render)
    lines.append('')

    if detail_messages:
        lines.append('## Appendix B — Message Details')
        lines.append('')
        for message in detail_messages:
            lines.append(f"### {message.get('label')} — {message.get('speaker_label')}")
            if message.get('timestamp'):
                lines.append(f"*{message.get('timestamp')}*")
            lines.append('')
            _append_markdown_mapping(lines, message.get('details', {}))
            lines.append('')

    if reference_messages:
        lines.append('## Appendix C — References')
        lines.append('')
        for message in reference_messages:
            lines.append(f"### {message.get('label')} — {message.get('speaker_label')}")
            if message.get('timestamp'):
                lines.append(f"*{message.get('timestamp')}*")
            lines.append('')
            _append_citations_markdown(lines, message)
            lines.append('')

    if thought_messages:
        lines.append('## Appendix D — Processing Thoughts')
        lines.append('')
        for message in thought_messages:
            lines.append(f"### {message.get('label')} — {message.get('speaker_label')}")
            if message.get('timestamp'):
                lines.append(f"*{message.get('timestamp')}*")
            lines.append('')
            for thought in message.get('thoughts', []):
                thought_label = thought.get('step_type', 'step').replace('_', ' ').title()
                lines.append(f"1. **{thought_label}:** {thought.get('content') or 'No content recorded.'}")
                if thought.get('duration_ms') is not None:
                    lines.append(f"   - **Duration:** {thought.get('duration_ms')} ms")
                if thought.get('timestamp'):
                    lines.append(f"   - **Timestamp:** {thought.get('timestamp')}")
                if thought.get('detail'):
                    lines.append('   - **Detail:**')
                    _append_code_block(lines, thought.get('detail'), indent='     ')
            lines.append('')

    if supplemental_messages:
        lines.append('## Appendix E — Supplemental Messages')
        lines.append('')
        for message in supplemental_messages:
            lines.append(f"### {message.get('label')} — {message.get('speaker_label')}")
            if message.get('timestamp'):
                lines.append(f"*{message.get('timestamp')}*")
            lines.append('')
            lines.append(message.get('content_text') or '_No content recorded._')
            lines.append('')

    return '\n'.join(lines).strip()


def _append_citations_markdown(lines: List[str], message: Dict[str, Any]):
    document_citations = [citation for citation in message.get('citations', []) if citation.get('citation_type') == 'document']
    web_citations = [citation for citation in message.get('citations', []) if citation.get('citation_type') == 'web']
    agent_citations = message.get('agent_citations', []) or []
    legacy_citations = [citation for citation in message.get('citations', []) if citation.get('citation_type') == 'legacy']

    if not any([document_citations, web_citations, agent_citations, legacy_citations]):
        lines.append('_No citations were recorded for this message._')
        return

    if document_citations:
        lines.append('#### Document Sources')
        lines.append('')
        for index, citation in enumerate(document_citations, start=1):
            lines.append(f"{index}. **{citation.get('label', 'Document source')}**")
            detail_mapping = _remove_empty_values({
                'citation_id': citation.get('citation_id'),
                'page_number': citation.get('page_number'),
                'classification': citation.get('classification'),
                'score': citation.get('score'),
                'metadata_type': citation.get('metadata_type')
            })
            _append_markdown_mapping(lines, detail_mapping, indent=1)
            if citation.get('metadata_content'):
                lines.append('   - **Metadata Content:**')
                _append_code_block(lines, citation.get('metadata_content'), indent='     ')
        lines.append('')

    if web_citations:
        lines.append('#### Web Sources')
        lines.append('')
        for index, citation in enumerate(web_citations, start=1):
            title = citation.get('title') or citation.get('label') or 'Web source'
            url = citation.get('url')
            if url:
                lines.append(f"{index}. [{title}]({url})")
            else:
                lines.append(f"{index}. {title}")
        lines.append('')

    if agent_citations:
        lines.append('#### Tool Invocations')
        lines.append('')
        for index, citation in enumerate(agent_citations, start=1):
            label = citation.get('tool_name') or citation.get('function_name') or f"Tool {index}"
            lines.append(f"{index}. **{label}**")
            detail_mapping = _remove_empty_values({
                'function_name': citation.get('function_name'),
                'plugin_name': citation.get('plugin_name'),
                'success': citation.get('success'),
                'timestamp': citation.get('timestamp')
            })
            _append_markdown_mapping(lines, detail_mapping, indent=1)
            if citation.get('function_arguments') not in (None, '', [], {}):
                lines.append('   - **Arguments:**')
                _append_code_block(lines, citation.get('function_arguments'), indent='     ')
            if citation.get('function_result') not in (None, '', [], {}):
                lines.append('   - **Result:**')
                _append_code_block(lines, citation.get('function_result'), indent='     ')
        lines.append('')

    if legacy_citations:
        lines.append('#### Legacy Citation Records')
        lines.append('')
        for index, citation in enumerate(legacy_citations, start=1):
            lines.append(f"{index}. {citation.get('label', 'Legacy citation')}")
        lines.append('')


def _append_markdown_mapping(lines: List[str], mapping: Dict[str, Any], indent: int = 0):
    if not isinstance(mapping, dict) or not mapping:
        return

    prefix = '  ' * indent
    for key, value in mapping.items():
        label = _format_markdown_key(key)
        if isinstance(value, dict):
            lines.append(f"{prefix}- **{label}:**")
            _append_markdown_mapping(lines, value, indent + 1)
        elif isinstance(value, list):
            if not value:
                continue
            if all(not isinstance(item, (dict, list)) for item in value):
                lines.append(f"{prefix}- **{label}:** {', '.join(_stringify_markdown_value(item) for item in value)}")
            else:
                lines.append(f"{prefix}- **{label}:**")
                for item in value:
                    if isinstance(item, dict):
                        lines.append(f"{prefix}  -")
                        _append_markdown_mapping(lines, item, indent + 2)
                    else:
                        lines.append(f"{prefix}  - {_stringify_markdown_value(item)}")
        else:
            lines.append(f"{prefix}- **{label}:** {_stringify_markdown_value(value)}")


def _append_code_block(lines: List[str], value: Any, indent: str = ''):
    if isinstance(value, (dict, list)):
        code_block = json.dumps(value, indent=2, ensure_ascii=False, default=str)
        language = 'json'
    else:
        code_block = str(value)
        language = 'text'

    lines.append(f"{indent}```{language}")
    for line in code_block.splitlines() or ['']:
        lines.append(f"{indent}{line}")
    lines.append(f"{indent}```")


def _format_markdown_key(key: str) -> str:
    return str(key).replace('_', ' ').title()


def _stringify_markdown_value(value: Any) -> str:
    if isinstance(value, bool):
        return 'Yes' if value else 'No'
    return str(value)


def _format_tag(tag: Any) -> str:
    """Format a tag or classification entry for display.

    Tags in Cosmos are stored as dicts such as
    ``{'category': 'model', 'value': 'gpt-5'}`` or
    ``{'category': 'participant', 'name': 'Alice', 'user_id': '...'}``
    but they can also be plain strings in older data.
    """
    if isinstance(tag, dict):
        category = tag.get('category', '')
        # Participant tags carry a readable name / email
        name = tag.get('name') or tag.get('email') or tag.get('display_name')
        if name:
            return f"{category}: {name}" if category else str(name)
        # Document tags carry a title
        title = tag.get('title') or tag.get('document_id')
        if title:
            return f"{category}: {title}" if category else str(title)
        # Generic category/value tags
        value = tag.get('value')
        if value:
            return f"{category}: {value}" if category else str(value)
        return category or str(tag)
    return str(tag)


def _role_to_label(role: str) -> str:
    role_map = {
        'assistant': 'Assistant',
        'user': 'User',
        'system': 'System',
        'tool': 'Tool',
        'file': 'File',
        'image': 'Image',
        'safety': 'Safety',
        'blocked': 'Blocked'
    }
    return role_map.get(role, str(role).capitalize() or 'Message')


def _normalize_content(content: Any) -> str:
    """Normalize message content to a plain string."""
    if isinstance(content, str):
        return content
    if isinstance(content, list):
        parts = []
        for item in content:
            if isinstance(item, dict):
                if item.get('type') == 'text':
                    parts.append(item.get('text', ''))
                elif item.get('type') == 'image_url':
                    parts.append('[Image]')
                else:
                    parts.append(str(item))
            else:
                parts.append(str(item))
        return '\n'.join(parts)
    if isinstance(content, dict):
        if content.get('type') == 'text':
            return content.get('text', '')
        return str(content)
    return str(content) if content else ''


def _safe_filename(title: str) -> str:
    """Create a filesystem-safe filename from a conversation title."""
    safe = re.sub(r'[<>:"/\\|?*]', '_', title)
    safe = re.sub(r'\s+', '_', safe)
    safe = safe.strip('_. ')
    if len(safe) > 50:
        safe = safe[:50]
    return safe or 'Untitled'


def _load_export_message_for_user(user_id: str, conversation_id: str, message_id: str) -> Dict[str, Any]:
    try:
        conversation = cosmos_conversations_container.read_item(
            item=conversation_id,
            partition_key=conversation_id
        )
    except Exception as exc:
        raise LookupError('Conversation not found') from exc

    if conversation.get('user_id') != user_id:
        raise PermissionError('Access denied')

    try:
        message = cosmos_messages_container.read_item(
            item=message_id,
            partition_key=conversation_id
        )
    except Exception:
        message_query = """
            SELECT * FROM c
            WHERE c.id = @message_id AND c.conversation_id = @conversation_id
        """
        message_results = list(cosmos_messages_container.query_items(
            query=message_query,
            parameters=[
                {'name': '@message_id', 'value': message_id},
                {'name': '@conversation_id', 'value': conversation_id}
            ],
            enable_cross_partition_query=True
        ))
        if not message_results:
            raise LookupError('Message not found')
        message = message_results[0]

    if message.get('conversation_id') != conversation_id:
        raise LookupError('Message not found')

    if isinstance(message.get('agent_citations'), list) and any(
        isinstance(citation, dict) and citation.get('artifact_id')
        for citation in message.get('agent_citations', [])
    ):
        conversation_messages = list(cosmos_messages_container.query_items(
            query="SELECT * FROM c WHERE c.conversation_id = @conversation_id",
            parameters=[{'name': '@conversation_id', 'value': conversation_id}],
            partition_key=conversation_id,
        ))
        artifact_payload_map = build_message_artifact_payload_map(conversation_messages)
        hydrated_messages = hydrate_agent_citations_from_artifacts([message], artifact_payload_map)
        if hydrated_messages:
            message = hydrated_messages[0]

    return message


def _message_to_docx_bytes(message: Dict[str, Any]) -> bytes:
    doc = DocxDocument()
    doc.add_heading('Message Export', level=1)

    role_label = _role_to_label(message.get('role', 'unknown'))
    timestamp = message.get('timestamp', '')

    meta_paragraph = doc.add_paragraph()
    meta_run = meta_paragraph.add_run(f"Role: {role_label}")
    meta_run.bold = True
    if timestamp:
        meta_paragraph.add_run(f"    {timestamp}")

    doc.add_paragraph('')

    content = _normalize_content(message.get('content', ''))
    if content:
        _add_markdown_content_to_doc(doc, content)
    else:
        doc.add_paragraph('No content recorded.')

    citation_labels = _build_message_citation_labels(message)
    if citation_labels:
        doc.add_heading('Citations', level=2)
        for citation_label in citation_labels:
            doc.add_paragraph(citation_label, style='List Bullet')

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()


def _message_to_email_draft_payload(
    message: Dict[str, Any],
    settings: Dict[str, Any],
    summary_model_deployment: str = ''
) -> Dict[str, Any]:
    content = _normalize_content(message.get('content', ''))
    subject_payload = _build_message_email_subject(
        content=content,
        settings=settings,
        requested_model=summary_model_deployment
    )
    body_content = _strip_explicit_message_email_subject(content)

    body_lines = []

    if body_content.strip():
        body_lines.extend(_render_markdown_to_email_lines(body_content))
    else:
        body_lines.append('No content recorded.')

    citation_labels = _build_message_citation_labels(message)
    if citation_labels:
        if body_lines and body_lines[-1] != '':
            body_lines.append('')
        body_lines.append('Citations')
        body_lines.append('---------')
        for citation_label in citation_labels:
            body_lines.append(f'- {citation_label}')

    body = _finalize_email_body_text(body_lines)
    return {
        'subject': subject_payload['subject'],
        'subject_source': subject_payload['source'],
        'body': body
    }


def _render_markdown_to_email_lines(content: str) -> List[str]:
    html = markdown2.markdown(content, extras=DOCX_MARKDOWN_EXTRAS)
    soup = BeautifulSoup(f'<div>{html}</div>', 'html.parser')
    root = soup.div if soup.div else soup
    lines: List[str] = []

    for child in root.children:
        if isinstance(child, NavigableString):
            text = str(child).strip()
            if text:
                lines.append(text)
                lines.append('')
            continue

        if isinstance(child, Tag):
            _append_html_block_to_email_lines(lines, child)

    return lines


def _append_html_block_to_email_lines(lines: List[str], node: Tag, list_level: int = 0):
    tag_name = node.name.lower()

    if tag_name in {'h1', 'h2', 'h3', 'h4', 'h5', 'h6'}:
        heading_text = _extract_email_inline_text(node).strip()
        if heading_text:
            underline_char = '=' if tag_name in {'h1', 'h2'} else '-'
            lines.append(heading_text)
            lines.append(underline_char * min(len(heading_text), 80))
            lines.append('')
        return

    if tag_name == 'p':
        paragraph_text = _extract_email_inline_text(node).strip()
        if paragraph_text:
            lines.extend(paragraph_text.splitlines())
            lines.append('')
        return

    if tag_name in {'ul', 'ol'}:
        _append_html_list_to_email_lines(lines, node, ordered=(tag_name == 'ol'), level=list_level)
        lines.append('')
        return

    if tag_name == 'pre':
        code_text = node.get_text().rstrip('\n')
        if code_text:
            for code_line in code_text.splitlines():
                lines.append(f'    {code_line.rstrip()}')
            lines.append('')
        return

    if tag_name == 'blockquote':
        quote_text = _extract_email_inline_text(node).strip()
        if quote_text:
            for quote_line in quote_text.splitlines():
                lines.append(f'    {quote_line}')
            lines.append('')
        return

    if tag_name == 'table':
        _append_html_table_to_email_lines(lines, node)
        lines.append('')
        return

    if tag_name == 'hr':
        lines.append('-' * 40)
        lines.append('')
        return

    if tag_name in {'div', 'section', 'article'}:
        for child in node.children:
            if isinstance(child, NavigableString):
                text = str(child).strip()
                if text:
                    lines.append(text)
                    lines.append('')
                continue

            if isinstance(child, Tag):
                _append_html_block_to_email_lines(lines, child, list_level=list_level)
        return

    fallback_text = _extract_email_inline_text(node).strip()
    if fallback_text:
        lines.extend(fallback_text.splitlines())
        lines.append('')


def _append_html_list_to_email_lines(lines: List[str], list_node: Tag, ordered: bool, level: int = 0):
    item_number = 1
    indent = '  ' * level

    for item in list_node.find_all('li', recursive=False):
        prefix = f'{item_number}. ' if ordered else '- '
        item_parts = []

        for child in item.children:
            if isinstance(child, Tag) and child.name.lower() in {'ul', 'ol'}:
                continue
            item_parts.append(_extract_email_inline_text(child))

        item_text = ''.join(item_parts).strip()
        if item_text:
            lines.append(f'{indent}{prefix}{item_text}')
        else:
            lines.append(f'{indent}{prefix}'.rstrip())

        for nested_list in item.find_all(['ul', 'ol'], recursive=False):
            _append_html_list_to_email_lines(
                lines,
                nested_list,
                ordered=(nested_list.name.lower() == 'ol'),
                level=level + 1
            )

        if ordered:
            item_number += 1


def _append_html_table_to_email_lines(lines: List[str], table_node: Tag):
    rows = table_node.find_all('tr')
    if not rows:
        return

    parsed_rows = []
    header_present = False
    for row_index, row in enumerate(rows):
        cells = row.find_all(['th', 'td'], recursive=False)
        if not cells:
            continue
        if row_index == 0 and all(cell.name.lower() == 'th' for cell in cells):
            header_present = True
        parsed_rows.append([
            re.sub(r'\s+', ' ', _extract_email_inline_text(cell)).strip()
            for cell in cells
        ])

    if not parsed_rows:
        return

    column_count = max(len(row) for row in parsed_rows)
    normalized_rows = [row + [''] * (column_count - len(row)) for row in parsed_rows]
    column_widths = [
        max(len(row[column_index]) for row in normalized_rows)
        for column_index in range(column_count)
    ]

    def format_row(row_values: List[str]) -> str:
        padded_cells = [
            row_values[column_index].ljust(column_widths[column_index])
            for column_index in range(column_count)
        ]
        return '  '.join(padded_cells).rstrip()

    lines.append(format_row(normalized_rows[0]))
    if header_present:
        separator = '  '.join(
            '-' * max(column_widths[column_index], 3)
            for column_index in range(column_count)
        )
        lines.append(separator)
        data_rows = normalized_rows[1:]
    else:
        data_rows = normalized_rows[1:]

    for row_values in data_rows:
        lines.append(format_row(row_values))


def _extract_email_inline_text(node: Any) -> str:
    if isinstance(node, NavigableString):
        return str(node)

    if not isinstance(node, Tag):
        return ''

    tag_name = node.name.lower()
    if tag_name == 'br':
        return '\n'
    if tag_name == 'img':
        return f"[{node.get('alt') or 'Image'}]"
    if tag_name == 'a':
        label = ''.join(_extract_email_inline_text(child) for child in node.children).strip()
        href = str(node.get('href') or '').strip()
        if href and href != label:
            if label:
                return f'{label} ({href})'
            return href
        return label

    return ''.join(_extract_email_inline_text(child) for child in node.children)


def _finalize_email_body_text(lines: List[str]) -> str:
    normalized_lines: List[str] = []

    for raw_line in lines:
        line = str(raw_line or '').rstrip()
        if not line:
            if normalized_lines and normalized_lines[-1] != '':
                normalized_lines.append('')
            continue
        normalized_lines.append(line)

    while normalized_lines and normalized_lines[-1] == '':
        normalized_lines.pop()

    return '\n'.join(normalized_lines)


def _build_message_email_subject(
    content: str,
    settings: Dict[str, Any],
    requested_model: str = ''
) -> Dict[str, str]:
    explicit_subject = _extract_message_email_subject(content)
    if explicit_subject:
        return {
            'subject': explicit_subject,
            'source': 'message'
        }

    generated_subject = _generate_message_email_subject_with_model(
        content=content,
        settings=settings,
        requested_model=requested_model
    )
    if generated_subject:
        return {
            'subject': generated_subject,
            'source': 'model'
        }

    return {
        'subject': _fallback_message_email_subject(content),
        'source': 'fallback'
    }


def _extract_message_email_subject(content: str) -> Optional[str]:
    if not content:
        return None

    lines = content.splitlines()
    explicit_patterns = [
        re.compile(r'^\s*(?:\*\*|__)?(?:email\s+)?subject(?:\*\*|__)?\s*:\s*(.+?)\s*$', re.IGNORECASE),
        re.compile(r'^\s*(?:\*\*|__)?title(?:\*\*|__)?\s*:\s*(.+?)\s*$', re.IGNORECASE),
    ]

    for line in lines:
        stripped_line = line.strip()
        if not stripped_line:
            continue
        for pattern in explicit_patterns:
            match = pattern.match(stripped_line)
            if not match:
                continue
            cleaned_subject = _clean_email_subject(match.group(1))
            if cleaned_subject:
                return cleaned_subject

    for line in lines:
        stripped_line = line.strip()
        if not stripped_line:
            continue

        heading_match = re.match(r'^#{1,6}\s+(.+)$', stripped_line)
        if heading_match:
            cleaned_subject = _clean_email_subject(heading_match.group(1))
            if cleaned_subject:
                return cleaned_subject
        break

    return None


def _strip_explicit_message_email_subject(content: str) -> str:
    if not content:
        return ''

    lines = content.splitlines()
    explicit_patterns = [
        re.compile(r'^\s*(?:\*\*|__)?(?:email\s+)?subject(?:\*\*|__)?\s*:\s*(.+?)\s*$', re.IGNORECASE),
        re.compile(r'^\s*(?:\*\*|__)?title(?:\*\*|__)?\s*:\s*(.+?)\s*$', re.IGNORECASE),
    ]

    first_non_empty_index = None
    for index, line in enumerate(lines):
        if line.strip():
            first_non_empty_index = index
            break

    if first_non_empty_index is None:
        return ''

    first_line = lines[first_non_empty_index].strip()
    if not any(pattern.match(first_line) for pattern in explicit_patterns):
        return content

    remaining_lines = lines[:first_non_empty_index] + lines[first_non_empty_index + 1:]
    while remaining_lines and not remaining_lines[0].strip():
        remaining_lines.pop(0)
    return '\n'.join(remaining_lines)


def _clean_email_subject(subject: str) -> str:
    cleaned_subject = re.sub(r'[`*_~]+', '', str(subject or ''))
    cleaned_subject = re.sub(r'\s+', ' ', cleaned_subject).strip()
    cleaned_subject = cleaned_subject.strip('"\'')
    cleaned_subject = cleaned_subject.rstrip(' .:;-')
    if len(cleaned_subject) > EMAIL_SUBJECT_CHAR_LIMIT:
        cleaned_subject = cleaned_subject[:EMAIL_SUBJECT_CHAR_LIMIT].rstrip(' .:;-')
    return cleaned_subject


def _generate_message_email_subject_with_model(
    content: str,
    settings: Dict[str, Any],
    requested_model: str = ''
) -> Optional[str]:
    subject_source = str(content or '').strip()
    if not subject_source:
        return None

    truncated_source = subject_source[:EMAIL_SUBJECT_SOURCE_CHAR_LIMIT]

    try:
        gpt_client, gpt_model = _initialize_gpt_client(settings, requested_model)
        model_lower = gpt_model.lower()
        is_reasoning_model = (
            'o1' in model_lower or 'o3' in model_lower or 'gpt-5' in model_lower
        )
        instruction_role = 'developer' if is_reasoning_model else 'system'
        subject_prompt = (
            'You are generating an email subject line for a mailto draft from a single chat message. '
            'If the message already contains a subject or clear title, reuse it in cleaned form. '
            'Otherwise, write a concise and specific subject line. '
            'Return plain text only with no quotes, no markdown, and no more than 10 words.'
        )

        subject_response = gpt_client.chat.completions.create(
            model=gpt_model,
            messages=[
                {
                    'role': instruction_role,
                    'content': subject_prompt
                },
                {
                    'role': 'user',
                    'content': truncated_source
                }
            ]
        )
        raw_subject = (
            (subject_response.choices[0].message.content or '').strip()
            if subject_response.choices else ''
        )
        cleaned_subject = _clean_email_subject(raw_subject)
        if cleaned_subject:
            return cleaned_subject
    except Exception as exc:
        debug_print(f'Message email subject generation failed: {exc}')
        log_event(
            'Message email subject generation failed',
            extra={
                'requested_model': requested_model or None,
                'content_length': len(subject_source)
            },
            level='WARNING'
        )

    return None


def _fallback_message_email_subject(content: str) -> str:
    extracted_subject = _extract_message_email_subject(content)
    if extracted_subject:
        return extracted_subject

    for line in str(content or '').splitlines():
        cleaned_subject = _clean_email_subject(line)
        if cleaned_subject:
            return cleaned_subject

    return 'Shared chat message'


def _build_message_citation_labels(message: Dict[str, Any]) -> List[str]:
    normalized_citations = _normalize_citations(_collect_raw_citation_buckets(message))
    citation_labels: List[str] = []
    seen_labels = set()

    for citation in normalized_citations:
        label = str(
            citation.get('label')
            or citation.get('title')
            or citation.get('url')
            or citation.get('filepath')
            or citation.get('tool_name')
            or citation.get('function_name')
            or ''
        ).strip()
        if not label or label in seen_labels:
            continue
        seen_labels.add(label)
        citation_labels.append(label)

    return citation_labels


def _add_markdown_content_to_doc(doc: DocxDocument, content: str):
    html = markdown2.markdown(content, extras=DOCX_MARKDOWN_EXTRAS)
    soup = BeautifulSoup(f'<div>{html}</div>', 'html.parser')
    root = soup.div if soup.div else soup
    rendered_blocks = False

    for child in root.children:
        if isinstance(child, NavigableString):
            text = str(child).strip()
            if not text:
                continue
            paragraph = doc.add_paragraph()
            paragraph.add_run(text)
            rendered_blocks = True
            continue

        if not isinstance(child, Tag):
            continue

        _append_html_block_to_doc(doc, child)
        rendered_blocks = True

    if not rendered_blocks and content.strip():
        doc.add_paragraph(content.strip())


def _append_html_block_to_doc(doc: DocxDocument, node: Tag, list_level: int = 0):
    tag_name = node.name.lower()

    if tag_name in {'h1', 'h2', 'h3', 'h4', 'h5', 'h6'}:
        paragraph = doc.add_heading('', level=min(int(tag_name[1]), 4))
        _append_inline_html_runs(paragraph, node)
        return

    if tag_name == 'p':
        paragraph = doc.add_paragraph()
        _append_inline_html_runs(paragraph, node)
        return

    if tag_name in {'ul', 'ol'}:
        _append_list_items_to_doc(doc, node, ordered=(tag_name == 'ol'), level=list_level)
        return

    if tag_name == 'pre':
        _add_code_block_to_doc(doc, node)
        return

    if tag_name == 'blockquote':
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.left_indent = Inches(0.3)
        _append_inline_html_runs(paragraph, node, {'italic': True})
        return

    if tag_name == 'table':
        _add_html_table_to_doc(doc, node)
        return

    if tag_name == 'hr':
        doc.add_paragraph('')
        return

    if tag_name in {'div', 'section', 'article'}:
        for child in node.children:
            if isinstance(child, NavigableString):
                text = str(child).strip()
                if not text:
                    continue
                paragraph = doc.add_paragraph()
                paragraph.add_run(text)
                continue

            if isinstance(child, Tag):
                _append_html_block_to_doc(doc, child, list_level=list_level)
        return

    paragraph = doc.add_paragraph()
    _append_inline_html_runs(paragraph, node)


def _append_list_items_to_doc(doc: DocxDocument, list_node: Tag, ordered: bool, level: int = 0):
    style_name = 'List Number' if ordered else 'List Bullet'

    for item in list_node.find_all('li', recursive=False):
        paragraph = doc.add_paragraph(style=style_name)
        if level:
            paragraph.paragraph_format.left_indent = Inches(0.25 * level)

        rendered_inline = False
        for child in item.children:
            if isinstance(child, Tag) and child.name.lower() in {'ul', 'ol'}:
                continue
            if isinstance(child, NavigableString) and not str(child).strip():
                continue

            _append_inline_html_runs(paragraph, child)
            rendered_inline = True

        if not rendered_inline:
            text = item.get_text(' ', strip=True)
            if text:
                paragraph.add_run(text)

        for nested_list in item.find_all(['ul', 'ol'], recursive=False):
            _append_list_items_to_doc(
                doc,
                nested_list,
                ordered=(nested_list.name.lower() == 'ol'),
                level=level + 1
            )


def _add_code_block_to_doc(doc: DocxDocument, node: Tag):
    code_text = node.get_text().rstrip('\n')
    if not code_text:
        return

    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.25)
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)


def _add_html_table_to_doc(doc: DocxDocument, table_node: Tag):
    rows = table_node.find_all('tr')
    if not rows:
        return

    column_count = max(
        len(row.find_all(['th', 'td'], recursive=False))
        for row in rows
    )
    if column_count == 0:
        return

    table = doc.add_table(rows=len(rows), cols=column_count)
    table.style = 'Table Grid'

    for row_index, row in enumerate(rows):
        cells = row.find_all(['th', 'td'], recursive=False)
        for column_index in range(column_count):
            cell = table.cell(row_index, column_index)
            cell.text = ''

            if column_index >= len(cells):
                continue

            _populate_table_cell(
                cell,
                cells[column_index],
                is_header=(cells[column_index].name.lower() == 'th')
            )


def _populate_table_cell(cell, node: Tag, is_header: bool = False):
    paragraph = cell.paragraphs[0]
    _append_inline_html_runs(paragraph, node, {'bold': is_header})


def _append_inline_html_runs(paragraph, node: Any, formatting: Optional[Dict[str, bool]] = None):
    if formatting is None:
        formatting = {}

    if isinstance(node, NavigableString):
        text = str(node)
        if not text:
            return

        run = paragraph.add_run(text)
        _apply_run_formatting(run, formatting)
        return

    if not isinstance(node, Tag):
        return

    tag_name = node.name.lower()
    if tag_name == 'br':
        paragraph.add_run().add_break()
        return

    if tag_name == 'img':
        alt_text = node.get('alt') or 'Image'
        run = paragraph.add_run(f'[{alt_text}]')
        _apply_run_formatting(run, formatting)
        return

    next_formatting = dict(formatting)
    if tag_name in {'strong', 'b'}:
        next_formatting['bold'] = True
    elif tag_name in {'em', 'i'}:
        next_formatting['italic'] = True
    elif tag_name in {'s', 'strike', 'del'}:
        next_formatting['strike'] = True
    elif tag_name == 'code':
        next_formatting['code'] = True
    elif tag_name == 'a':
        next_formatting['underline'] = True

    for child in node.children:
        _append_inline_html_runs(paragraph, child, next_formatting)

    if tag_name == 'a':
        href = str(node.get('href') or '').strip()
        label = node.get_text(' ', strip=True)
        if href and href != label:
            suffix_run = paragraph.add_run(f' ({href})')
            _apply_run_formatting(suffix_run, formatting)


def _apply_run_formatting(run, formatting: Dict[str, bool]):
    if formatting.get('bold'):
        run.bold = True
    if formatting.get('italic'):
        run.italic = True
    if formatting.get('underline'):
        run.underline = True
    if formatting.get('strike'):
        run.font.strike = True
    if formatting.get('code'):
        run.font.name = 'Consolas'
        run.font.size = Pt(9)


# ---------------------------------------------------------------------------
# PDF Export — HTML generation and PyMuPDF Story rendering
# ---------------------------------------------------------------------------

_PDF_CSS = """
body {
    font-family: sans-serif;
    font-size: 10pt;
    color: #222;
    line-height: 1.4;
}
h1 {
    font-size: 16pt;
    color: #1a1a2e;
    margin-bottom: 2pt;
}
h2 {
    font-size: 13pt;
    color: #16213e;
    margin-top: 16pt;
    margin-bottom: 6pt;
    border-bottom: 1px solid #ccc;
    padding-bottom: 4pt;
}
h3 {
    font-size: 11pt;
    color: #0f3460;
    margin-top: 10pt;
    margin-bottom: 4pt;
}
h4 {
    font-size: 10pt;
    color: #333;
    margin-top: 8pt;
    margin-bottom: 4pt;
}
p {
    margin-top: 2pt;
    margin-bottom: 4pt;
}
.metadata {
    font-size: 8pt;
    color: #666;
}
.abstract {
    background-color: #f8f9fa;
    padding: 8pt;
    margin-bottom: 8pt;
}
.note {
    font-size: 9pt;
    color: #856404;
    background-color: #fff3cd;
    padding: 6pt;
}
.bubble {
    padding: 8pt 12pt;
    margin-bottom: 8pt;
}
.bubble-header {
    font-size: 8pt;
    color: #444;
    margin-bottom: 2pt;
}
.ts {
    font-weight: normal;
    color: #888;
}
.user-bubble {
    background-color: #c8e0fa;
    margin-left: 60pt;
}
.assistant-bubble {
    background-color: #f1f0f0;
    margin-right: 60pt;
}
.system-bubble {
    background-color: #fff3cd;
    margin-left: 30pt;
    margin-right: 30pt;
    font-size: 9pt;
}
.file-bubble {
    background-color: #e8f5e9;
    margin-right: 60pt;
    font-size: 9pt;
}
.other-bubble {
    background-color: #f5f5f5;
    margin-left: 30pt;
    margin-right: 30pt;
    font-size: 9pt;
}
table {
    border-collapse: collapse;
    width: 100%;
    font-size: 9pt;
    margin-bottom: 8pt;
}
th, td {
    border: 1px solid #ddd;
    padding: 4pt 6pt;
    text-align: left;
}
th {
    background-color: #f5f5f5;
    font-weight: bold;
}
pre {
    background-color: #f5f5f5;
    padding: 6pt;
    font-size: 8pt;
    font-family: monospace;
}
code {
    font-family: monospace;
    font-size: 9pt;
    background-color: #f0f0f0;
    padding: 1pt 3pt;
}
ol, ul {
    margin-top: 4pt;
    margin-bottom: 8pt;
}
li {
    margin-bottom: 4pt;
}
small {
    font-size: 8pt;
    color: #666;
}
a {
    color: #0066cc;
}
"""


def _pdf_bubble_class(role: str) -> str:
    """Return the CSS class for a chat bubble based on message role."""
    role_classes = {
        'user': 'user-bubble',
        'assistant': 'assistant-bubble',
        'system': 'system-bubble',
        'file': 'file-bubble',
        'image': 'file-bubble'
    }
    return role_classes.get(role, 'other-bubble')


def _build_pdf_html_body(entry: Dict[str, Any]) -> str:
    """Build the HTML body content for a single conversation PDF."""
    conversation = entry['conversation']
    messages = entry['messages']
    summary_intro = entry.get('summary_intro', {}) or {}

    transcript_messages = [m for m in messages if m.get('is_transcript_message')]
    detail_messages = [m for m in messages if m.get('details')]
    reference_messages = [m for m in messages if m.get('citations')]
    thought_messages = [m for m in messages if m.get('thoughts')]
    supplemental_messages = [m for m in messages if not m.get('is_transcript_message')]

    parts: List[str] = []

    # --- Title and metadata ---
    parts.append(f'<h1>{_escape_html(conversation.get("title", "Untitled"))}</h1>')
    meta_items = [
        f'<b>Last Updated:</b> {_escape_html(str(conversation.get("last_updated", "")))}',
        f'<b>Chat Type:</b> {_escape_html(str(conversation.get("chat_type", "personal")))}',
        f'<b>Messages:</b> {conversation.get("message_count", len(messages))}'
    ]
    tags = conversation.get('tags')
    if tags:
        meta_items.append(f'<b>Tags:</b> {_escape_html(", ".join(_format_tag(t) for t in tags))}')
    classification = conversation.get('classification')
    if classification:
        meta_items.append(
            f'<b>Classification:</b> {_escape_html(", ".join(_format_tag(c) for c in classification))}'
        )
    parts.append(f'<p class="metadata">{" &nbsp;|&nbsp; ".join(meta_items)}</p>')

    # --- Abstract ---
    if summary_intro.get('enabled') and summary_intro.get('generated') and summary_intro.get('content'):
        parts.append('<h2>Abstract</h2>')
        abstract_html = markdown2.markdown(
            summary_intro.get('content', ''),
            extras=['fenced-code-blocks', 'tables']
        )
        parts.append(f'<div class="abstract">{abstract_html}</div>')
        parts.append(
            f'<p class="metadata"><i>Generated with '
            f'{_escape_html(str(summary_intro.get("model_deployment") or "configured model"))} on '
            f'{_escape_html(str(summary_intro.get("generated_at", "")))}</i></p>'
        )
    elif summary_intro.get('enabled') and summary_intro.get('error'):
        error_text = _escape_html(str(summary_intro.get('error', '')))
        parts.append(
            '<p class="note"><i>A summary intro was requested, '
            'but could not be generated for this export.</i><br/>'
            f'<small>Error: {error_text}</small></p>'
        )

    # --- Transcript with chat bubbles ---
    parts.append('<h2>Transcript</h2>')
    if not transcript_messages:
        parts.append(
            '<p><i>No user or assistant transcript messages were available for export.</i></p>'
        )
    else:
        for message in transcript_messages:
            role = message.get('role', '')
            bubble_class = _pdf_bubble_class(role)
            label = message.get('label', '')
            speaker = message.get('speaker_label', '')
            timestamp = message.get('timestamp', '')
            content = message.get('content_text', '') or 'No content recorded.'

            parts.append(f'<div class="bubble {bubble_class}">')
            ts_str = (
                f' &nbsp;|&nbsp; <span class="ts">{_escape_html(str(timestamp))}</span>'
                if timestamp else ''
            )
            parts.append(
                f'<p class="bubble-header"><b>{_escape_html(label)} — '
                f'{_escape_html(speaker)}</b>{ts_str}</p>'
            )
            content_html = markdown2.markdown(
                content,
                extras=['fenced-code-blocks', 'tables', 'break-on-newline']
            )
            parts.append(content_html)
            parts.append('</div>')

    # --- Appendix A: Conversation Metadata ---
    parts.append('<h2>Appendix A — Conversation Metadata</h2>')
    metadata_to_render = _remove_empty_values({
        'context': conversation.get('context'),
        'classification': conversation.get('classification'),
        'strict': conversation.get('strict'),
        'is_pinned': conversation.get('is_pinned'),
        'scope_locked': conversation.get('scope_locked'),
        'locked_contexts': conversation.get('locked_contexts'),
        'message_counts_by_role': conversation.get('message_counts_by_role'),
        'citation_counts': conversation.get('citation_counts'),
        'thought_count': conversation.get('thought_count')
    })
    _append_html_table(parts, metadata_to_render)

    # --- Appendix B: Message Details ---
    if detail_messages:
        parts.append('<h2>Appendix B — Message Details</h2>')
        for message in detail_messages:
            parts.append(
                f'<h3>{_escape_html(message.get("label", ""))} — '
                f'{_escape_html(message.get("speaker_label", ""))}</h3>'
            )
            if message.get('timestamp'):
                parts.append(
                    f'<p class="metadata"><i>{_escape_html(str(message.get("timestamp")))}</i></p>'
                )
            _append_html_table(parts, message.get('details', {}))

    # --- Appendix C: References ---
    if reference_messages:
        parts.append('<h2>Appendix C — References</h2>')
        for message in reference_messages:
            parts.append(
                f'<h3>{_escape_html(message.get("label", ""))} — '
                f'{_escape_html(message.get("speaker_label", ""))}</h3>'
            )
            if message.get('timestamp'):
                parts.append(
                    f'<p class="metadata"><i>{_escape_html(str(message.get("timestamp")))}</i></p>'
                )
            _append_html_citations(parts, message)

    # --- Appendix D: Processing Thoughts ---
    if thought_messages:
        parts.append('<h2>Appendix D — Processing Thoughts</h2>')
        for message in thought_messages:
            parts.append(
                f'<h3>{_escape_html(message.get("label", ""))} — '
                f'{_escape_html(message.get("speaker_label", ""))}</h3>'
            )
            if message.get('timestamp'):
                parts.append(
                    f'<p class="metadata"><i>{_escape_html(str(message.get("timestamp")))}</i></p>'
                )
            parts.append('<ol>')
            for thought in message.get('thoughts', []):
                thought_label = (thought.get('step_type') or 'step').replace('_', ' ').title()
                parts.append(
                    f'<li><b>{_escape_html(thought_label)}:</b> '
                    f'{_escape_html(str(thought.get("content") or "No content recorded."))}'
                )
                if thought.get('duration_ms') is not None:
                    parts.append(
                        f'<br/><small><b>Duration:</b> {thought.get("duration_ms")} ms</small>'
                    )
                if thought.get('timestamp'):
                    parts.append(
                        f'<br/><small><b>Timestamp:</b> '
                        f'{_escape_html(str(thought.get("timestamp")))}</small>'
                    )
                if thought.get('detail'):
                    parts.append('<br/><small><b>Detail:</b></small>')
                    _append_html_code_block(parts, thought.get('detail'))
                parts.append('</li>')
            parts.append('</ol>')

    # --- Appendix E: Supplemental Messages ---
    if supplemental_messages:
        parts.append('<h2>Appendix E — Supplemental Messages</h2>')
        for message in supplemental_messages:
            parts.append(
                f'<h3>{_escape_html(message.get("label", ""))} — '
                f'{_escape_html(message.get("speaker_label", ""))}</h3>'
            )
            if message.get('timestamp'):
                parts.append(
                    f'<p class="metadata"><i>{_escape_html(str(message.get("timestamp")))}</i></p>'
                )
            content = message.get('content_text', '') or 'No content recorded.'
            content_html = markdown2.markdown(
                content,
                extras=['fenced-code-blocks', 'tables', 'break-on-newline']
            )
            parts.append(content_html)

    return '\n'.join(parts)


def _render_pdf_bytes(body_html: str) -> bytes:
    """Render HTML body content to PDF bytes using PyMuPDF Story API."""
    MEDIABOX = fitz.paper_rect("letter")
    WHERE = MEDIABOX + (36, 36, -36, -36)

    story = fitz.Story(html=body_html, user_css=_PDF_CSS)

    tmp_path = None
    try:
        with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as tmp:
            tmp_path = tmp.name

        writer = fitz.DocumentWriter(tmp_path)
        more = True
        while more:
            device = writer.begin_page(MEDIABOX)
            more, _ = story.place(WHERE)
            story.draw(device)
            writer.end_page()
        writer.close()
        del story
        del writer

        with open(tmp_path, 'rb') as f:
            return f.read()
    finally:
        if tmp_path:
            try:
                os.unlink(tmp_path)
            except OSError:
                pass


def _conversation_to_pdf_bytes(entry: Dict[str, Any]) -> bytes:
    """Convert a conversation export entry to PDF bytes."""
    body_html = _build_pdf_html_body(entry)
    return _render_pdf_bytes(body_html)


def _html_body_to_pdf_bytes(body_html: str) -> bytes:
    """Convert raw HTML body content to PDF bytes."""
    return _render_pdf_bytes(body_html)


def _append_html_table(parts: List[str], mapping: Dict[str, Any]):
    """Append a key-value mapping as an HTML table."""
    if not isinstance(mapping, dict) or not mapping:
        parts.append('<p><i>No data available.</i></p>')
        return

    parts.append('<table>')
    parts.append('<tr><th>Property</th><th>Value</th></tr>')
    for key, value in mapping.items():
        label = _format_markdown_key(key)
        if isinstance(value, dict):
            formatted = _format_nested_html_value(value)
        elif isinstance(value, list):
            formatted = (
                ', '.join(_escape_html(str(item)) for item in value)
                if value else '<i>None</i>'
            )
        elif isinstance(value, bool):
            formatted = 'Yes' if value else 'No'
        else:
            formatted = _escape_html(str(value))
        parts.append(f'<tr><td><b>{_escape_html(label)}</b></td><td>{formatted}</td></tr>')
    parts.append('</table>')


def _format_nested_html_value(mapping: Dict[str, Any], depth: int = 0) -> str:
    """Format a nested dict as an HTML string for table cells."""
    if not mapping:
        return '<i>None</i>'

    items = []
    for key, value in mapping.items():
        label = _format_markdown_key(key)
        if isinstance(value, dict):
            nested = _format_nested_html_value(value, depth + 1)
            items.append(f'<b>{_escape_html(label)}:</b><br/>{nested}')
        elif isinstance(value, list):
            list_str = (
                ', '.join(_escape_html(str(v)) for v in value)
                if value else 'None'
            )
            items.append(f'<b>{_escape_html(label)}:</b> {list_str}')
        elif isinstance(value, bool):
            items.append(f'<b>{_escape_html(label)}:</b> {"Yes" if value else "No"}')
        else:
            items.append(f'<b>{_escape_html(label)}:</b> {_escape_html(str(value))}')
    return '<br/>'.join(items)


def _append_html_citations(parts: List[str], message: Dict[str, Any]):
    """Append citation data as HTML."""
    citations = message.get('citations', [])
    if not citations:
        parts.append('<p><i>No citations were recorded for this message.</i></p>')
        return

    doc_citations = [c for c in citations if c.get('citation_type') == 'document']
    web_citations = [c for c in citations if c.get('citation_type') == 'web']
    agent_citations = [c for c in citations if c.get('citation_type') == 'agent_tool']
    legacy_citations = [c for c in citations if c.get('citation_type') == 'legacy']

    if doc_citations:
        parts.append('<h4>Document Sources</h4>')
        parts.append('<ol>')
        for citation in doc_citations:
            parts.append(
                f'<li><b>{_escape_html(str(citation.get("label", "Document source")))}</b>'
            )
            detail_items = _remove_empty_values({
                'citation_id': citation.get('citation_id'),
                'page_number': citation.get('page_number'),
                'classification': citation.get('classification'),
                'score': citation.get('score'),
                'metadata_type': citation.get('metadata_type')
            })
            if detail_items:
                detail_str = '; '.join(
                    f'{_format_markdown_key(k)}: {_escape_html(str(v))}'
                    for k, v in detail_items.items()
                )
                parts.append(f'<br/><small>{detail_str}</small>')
            if citation.get('metadata_content'):
                parts.append('<br/><small><b>Metadata Content:</b></small>')
                _append_html_code_block(parts, citation.get('metadata_content'))
            parts.append('</li>')
        parts.append('</ol>')

    if web_citations:
        parts.append('<h4>Web Sources</h4>')
        parts.append('<ol>')
        for citation in web_citations:
            title = _escape_html(
                str(citation.get('title') or citation.get('label') or 'Web source')
            )
            url = citation.get('url')
            if url:
                parts.append(f'<li><a href="{_escape_html(url)}">{title}</a></li>')
            else:
                parts.append(f'<li>{title}</li>')
        parts.append('</ol>')

    if agent_citations:
        parts.append('<h4>Tool Invocations</h4>')
        parts.append('<ol>')
        for citation in agent_citations:
            label = _escape_html(
                str(citation.get('tool_name') or citation.get('function_name') or 'Tool')
            )
            parts.append(f'<li><b>{label}</b>')
            detail_items = _remove_empty_values({
                'function_name': citation.get('function_name'),
                'plugin_name': citation.get('plugin_name'),
                'success': citation.get('success'),
                'timestamp': citation.get('timestamp')
            })
            if detail_items:
                detail_str = '; '.join(
                    f'{_format_markdown_key(k)}: {_escape_html(str(v))}'
                    for k, v in detail_items.items()
                )
                parts.append(f'<br/><small>{detail_str}</small>')
            parts.append('</li>')
        parts.append('</ol>')

    if legacy_citations:
        parts.append('<h4>Legacy Citation Records</h4>')
        parts.append('<ol>')
        for citation in legacy_citations:
            parts.append(
                f'<li>{_escape_html(str(citation.get("label", "Legacy citation")))}</li>'
            )
        parts.append('</ol>')


def _append_html_code_block(parts: List[str], value: Any):
    """Append a code block in HTML format."""
    if isinstance(value, (dict, list)):
        code_text = json.dumps(value, indent=2, ensure_ascii=False, default=str)
    else:
        code_text = str(value)
    parts.append(f'<pre>{_escape_html(code_text)}</pre>')

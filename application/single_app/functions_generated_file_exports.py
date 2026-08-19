# functions_generated_file_exports.py
"""Format-neutral planning and rendering for generated chat file exports."""

import html
import io
import json
import os
import re
import tempfile
from datetime import datetime
from typing import Any, Dict, Iterable, List, Optional, Sequence, Tuple
from xml.etree import ElementTree

from defusedxml import ElementTree as DefusedElementTree
from defusedxml.common import DefusedXmlException

from functions_assistant_table_exports import (
    assistant_table_export_requested,
    build_assistant_table_csv,
    build_csv_output_clarification_guidance,
    extract_assistant_table_entries,
)


GENERATED_FILE_FORMAT_CSV = 'csv'
GENERATED_FILE_FORMAT_DOCX = 'docx'
GENERATED_FILE_FORMAT_PDF = 'pdf'
GENERATED_FILE_FORMATS = {
    GENERATED_FILE_FORMAT_CSV,
    GENERATED_FILE_FORMAT_DOCX,
    GENERATED_FILE_FORMAT_PDF,
}
SUPPORTED_GENERATED_EXPORT_FORMATS = {'csv', 'json', 'xml'}
# Only these payloads are withheld from the streamed assistant text, so only their cards
# replace it. CSV narratives stream intact and must stay visible above the artifact card.
ASSISTANT_TEXT_SUPPRESSING_FORMATS = {'json', 'xml'}
GENERATED_FILE_PREVIEW_ROWS = 3
REQUESTED_ARTIFACT_FORMATS = ('csv', 'json', 'xml', 'md', 'docx', 'pdf')
STRUCTURED_ARTIFACT_FORMAT_MARKERS = {
    'json': (
        'json artifact',
        'json export',
        'json output',
        'convert into json',
        'convert to json',
        'return json',
        'return only json',
        'respond with json',
        'format as json',
        'output as json',
        'save as json',
        'export as json',
        'download as json',
        'create json',
        'create a json',
        'make json',
        'make a json',
        'generate json',
        'generate a json',
    ),
    'xml': (
        'xml artifact',
        'xml export',
        'xml output',
        'convert into xml',
        'convert to xml',
        'populate xml',
        'populate the xml',
        'return xml',
        'return only xml',
        'respond with xml',
        'format as xml',
        'output as xml',
        'save as xml',
        'export as xml',
        'download as xml',
        'create xml',
        'create an xml',
        'make xml',
        'make an xml',
        'generate xml',
        'generate an xml',
    ),
}
STRUCTURED_ARTIFACT_ACTION_PATTERN = (
    r'convert|populate|create|make|build|generate|produce|return|respond|format|output|save|export|download'
)
STRUCTURED_ARTIFACT_DESTINATION_ACTION_PATTERN = (
    r'convert|transform|translate|turn|put|place|write|map|load|insert|transfer|copy|move'
)
FUNCTION_RESULT_ROW_KEYS = (
    'rows',
    'data',
    'items',
    'results',
    'records',
    'value',
    'values',
    'result',
    'body',
    'output',
    'payload',
)
FUNCTION_RESULT_CONTROL_KEYS = {
    'count',
    'detail',
    'error',
    'errormessage',
    'hasmore',
    'message',
    'metadata',
    'meta',
    'nextlink',
    'nextpage',
    'pagination',
    'returnedrows',
    'status',
    'statuscode',
    'success',
    'summary',
    'total',
    'totalcount',
    'totalmatches',
}
FUNCTION_RESULT_SENSITIVE_KEY_FRAGMENTS = (
    'accesstoken',
    'apikey',
    'authorization',
    'clientsecret',
    'connectionstring',
    'credential',
    'password',
    'privatekey',
    'secret',
    'sharedaccesssignature',
    'subscriptionkey',
    'token',
)
TABULAR_FUNCTION_RESULT_PLUGIN_NAMES = {'tabularprocessingplugin'}
XML_DECLARATION = '<?xml version="1.0" encoding="UTF-8"?>'
XML_ROOT_PATTERN = re.compile(r'<(?P<tag>[A-Za-z_][A-Za-z0-9_.:-]*)(?:\s[^<>]*)?>')
DOCX_OUTPUT_REQUEST_PATTERNS = (
    re.compile(
        r'\b(?:build|create|download|export|generate|make|prepare|save|turn|convert)\b'
        r'.{0,120}\b(?:a\s+)?(?:word|docx)(?:\s+(?:document|file|output|report))?\b'
    ),
    re.compile(r'\b(?:word|docx)\s+(?:document|file|output|report|version)\b'),
    re.compile(r'\b(?:get|give)\s+(?:me\s+)?(?:a|the|one)\s+(?:word|docx)\b'),
    re.compile(r'\b(?:need|want)\s+(?:(?:a|the|one)\s+)?(?:word|docx)\b'),
    re.compile(r'\b(?:in|as)\s+(?:a\s+)?(?:word|docx)(?:\s+(?:document|file|report))?\b'),
)
PDF_OUTPUT_REQUEST_PATTERNS = (
    re.compile(
        r'\b(?:build|create|download|export|generate|make|prepare|save|turn|convert)\b'
        r'.{0,120}\b(?:a\s+)?pdf(?:\s+(?:document|file|output|report))?\b'
    ),
    re.compile(r'\bpdf\s+(?:document|file|output|report|version)\b'),
    re.compile(r'\b(?:get|give)\s+(?:me\s+)?(?:a|the|one)\s+pdf\b'),
    re.compile(r'\b(?:need|want)\s+(?:(?:a|the|one)\s+)?pdf\b'),
    re.compile(r'\b(?:in|as)\s+(?:a\s+)?pdf(?:\s+(?:document|file|report))?\b'),
)
PDF_EXPORT_CSS = """
body { font-family: sans-serif; font-size: 10pt; color: #172033; }
h1 { font-size: 20pt; color: #173b5f; margin-bottom: 10pt; }
h2 { font-size: 14pt; color: #173b5f; margin-top: 16pt; }
p { line-height: 1.35; margin-bottom: 8pt; }
table { border-collapse: collapse; width: 100%; margin-top: 8pt; }
th { background-color: #e8eef5; font-weight: bold; }
th, td { border: 0.6pt solid #aab7c4; padding: 4pt; vertical-align: top; }
"""

MARKDOWN_OUTPUT_REQUEST_PATTERNS = (
    re.compile(
        r'\b(?:build|create|download|export|generate|make|prepare|save|write)\b'
        r'.{0,120}\b(?:markdown|md)(?:\s+(?:analysis|artifact|document|file|output|report))?\b'
    ),
    re.compile(r'\b(?:markdown|md)\s+(?:analysis|artifact|document|file|output|report|version)\b'),
    re.compile(r'\b(?:in|as)\s+(?:a\s+)?(?:markdown|md)(?:\s+(?:document|file|report))?\b'),
)

PASSTHROUGH_DERIVED_OUTPUT_PATTERNS = (
    re.compile(r'\b(?:derive|derived|classify|classification|categorize|category|calculate|computed?|map|mapping)\b'),
    re.compile(r'\b(?:score|rank|judge|evaluate|determine|flag|label|extract|populate|fill)\b'),
    re.compile(r'\b(?:analy[sz]e|summari[sz]e)\b'),
    re.compile(r'\b(?:exactly|only)\s+(?:these\s+)?(?:fields|columns)\b'),
    re.compile(r'\b(?:output|requested|derived)\s+(?:fields|columns|schema)\b'),
    re.compile(r'\bone\s+output\s+row\s+(?:for|per)\s+(?:each|every|source)\s+row\b'),
)
PASSTHROUGH_COPY_PATTERNS = (
    re.compile(r'\b(?:unchanged|as-is|as\s+is|verbatim|raw|original)\s+(?:copy|rows?|data|table|result|results)\b'),
    re.compile(r'\b(?:copy|export|download|save)\b[\w\s,.:;\-/]{0,100}\b(?:unchanged|as-is|as\s+is|verbatim|raw|original|source)\b'),
)
PASSTHROUGH_SERIALIZE_PATTERNS = (
    re.compile(r'\b(?:build|create|download|export|format|generate|make|prepare|save|serialize|convert)\b[\w\s,.:;\-/]{0,100}\b(?:csv|json|xml|docx|word|pdf|spreadsheet)\b'),
    re.compile(r'\b(?:csv|json|xml|docx|word|pdf|spreadsheet)\b[\w\s,.:;\-/]{0,100}\b(?:export|download|file|format|copy)\b'),
)


def _normalize_question_for_artifact_detection(user_question: str) -> str:
    return re.sub(r'\s+', ' ', str(user_question or '').strip().casefold())


def _iter_request_clauses(normalized_question: str):
    for match in re.finditer(r'[^.!?;\n]+', normalized_question):
        clause = match.group(0).strip()
        if clause:
            leading_offset = len(match.group(0)) - len(match.group(0).lstrip())
            yield clause, match.start() + leading_offset


def _first_pattern_position(normalized_question: str, patterns) -> Optional[int]:
    positions = [match.start() for pattern in patterns for match in [pattern.search(normalized_question)] if match]
    return min(positions) if positions else None


def _format_aliases(output_format: str) -> Tuple[str, ...]:
    aliases = {
        'docx': ('docx', 'word'),
        'md': ('md', 'markdown'),
    }
    return aliases.get(output_format, (output_format,))


def _clause_negates_output_format(clause: str, output_format: str) -> bool:
    return any(
        _structured_artifact_format_is_negated(clause, format_alias)
        for format_alias in _format_aliases(output_format)
    )


def _first_csv_artifact_position(user_question: str, normalized_question: str) -> Optional[int]:
    if not assistant_table_export_requested(user_question):
        return None
    for clause, clause_offset in _iter_request_clauses(normalized_question):
        if _clause_negates_output_format(clause, 'csv'):
            continue
        for marker in ('csv', 'spreadsheet'):
            marker_position = clause.find(marker)
            if marker_position >= 0:
                return clause_offset + marker_position
    return 0


def _collect_structured_artifact_format_matches(normalized_question: str) -> List[Tuple[int, str]]:
    matches = []
    destination_formats_by_clause_offset = {}
    for clause, clause_offset in _iter_request_clauses(normalized_question):
        for output_format in ('json', 'xml'):
            if output_format not in clause or _clause_negates_output_format(clause, output_format):
                continue
            destination_match = re.search(
                rf'\b(?:{STRUCTURED_ARTIFACT_DESTINATION_ACTION_PATTERN})\b'
                rf'[\w\s,():;\-/]{{0,120}}\b(?:into|in|to|as)\s+'
                rf'(?:(?:an?|the)\s+)?(?:new\s+)?{output_format}'
                rf'(?:\s+(?:artifact|document|file|format|output|template))?\b',
                clause,
            )
            if destination_match:
                format_position = clause.find(output_format, destination_match.start())
                matches.append((clause_offset + (format_position if format_position >= 0 else destination_match.start()), output_format))
                destination_formats_by_clause_offset.setdefault(clause_offset, set()).add(output_format)

    for clause, clause_offset in _iter_request_clauses(normalized_question):
        destination_formats = destination_formats_by_clause_offset.get(clause_offset, set())
        for output_format in ('json', 'xml'):
            if output_format not in clause or _clause_negates_output_format(clause, output_format):
                continue
            if destination_formats and output_format not in destination_formats:
                continue
            marker_positions = [
                clause.find(marker)
                for marker in STRUCTURED_ARTIFACT_FORMAT_MARKERS[output_format]
                if marker in clause
            ]
            if marker_positions:
                matches.append((clause_offset + min(marker_positions), output_format))
                continue
            generic_match = re.search(
                rf'\b(?:{STRUCTURED_ARTIFACT_ACTION_PATTERN})\b'
                rf'[\w\s.,:;\-/]{{0,80}}\b(?:an?\s+)?{output_format}\b',
                clause,
            )
            if generic_match:
                format_position = clause.find(output_format, generic_match.start())
                matches.append((clause_offset + (format_position if format_position >= 0 else generic_match.start()), output_format))
    return matches


def get_requested_artifact_formats(user_question: str) -> List[str]:
    """Return explicitly requested artifact formats in user-request order."""
    normalized_question = _normalize_question_for_artifact_detection(user_question)
    if not normalized_question:
        return []

    matches = []
    csv_position = _first_csv_artifact_position(user_question, normalized_question)
    if csv_position is not None:
        matches.append((csv_position, 'csv'))
    matches.extend(_collect_structured_artifact_format_matches(normalized_question))

    format_pattern_sets = {
        'md': MARKDOWN_OUTPUT_REQUEST_PATTERNS,
        'docx': DOCX_OUTPUT_REQUEST_PATTERNS,
        'pdf': PDF_OUTPUT_REQUEST_PATTERNS,
    }
    for output_format, patterns in format_pattern_sets.items():
        position = _first_pattern_position(normalized_question, patterns)
        if position is None:
            continue
        containing_clause = next(
            (
                clause
                for clause, clause_offset in _iter_request_clauses(normalized_question)
                if clause_offset <= position < clause_offset + len(clause)
            ),
            normalized_question,
        )
        if _clause_negates_output_format(containing_clause, output_format):
            continue
        matches.append((position, output_format))

    ordered_formats = []
    for _, output_format in sorted(matches, key=lambda item: (item[0], REQUESTED_ARTIFACT_FORMATS.index(item[1]))):
        if output_format not in ordered_formats:
            ordered_formats.append(output_format)
    return ordered_formats


def get_requested_structured_artifact_formats(user_question: str) -> List[str]:
    """Return requested durable structured artifact formats in user-request order."""
    return [
        output_format
        for output_format in get_requested_artifact_formats(user_question)
        if output_format in SUPPORTED_GENERATED_EXPORT_FORMATS
    ]


def get_requested_generated_file_formats(user_question: str) -> List[str]:
    """Return requested single-reply generated file formats in user-request order."""
    return [
        output_format
        for output_format in get_requested_artifact_formats(user_question)
        if output_format in GENERATED_FILE_FORMATS
    ]


def get_requested_generated_file_format(user_question: str) -> Optional[str]:
    """Return the requested generated file format, if any."""
    requested_formats = get_requested_generated_file_formats(user_question)
    return requested_formats[0] if requested_formats else None


def get_requested_structured_artifact_format(user_question: str) -> Optional[str]:
    """Return a requested CSV, JSON, or XML artifact target without resolving source orchestration."""
    requested_formats = get_requested_structured_artifact_formats(user_question)
    return requested_formats[0] if requested_formats else None


def _structured_artifact_format_is_negated(clause: str, output_format: str) -> bool:
    """Return whether a clause directly negates creating the target structured format."""
    return bool(
        re.search(
            rf"\b(?:do\s+not|don't|dont|never)\s+"
            rf"(?:{STRUCTURED_ARTIFACT_ACTION_PATTERN}|{STRUCTURED_ARTIFACT_DESTINATION_ACTION_PATTERN})"
            rf"\b[\w\s.,:;\-/]{{0,80}}\b{output_format}\b",
            clause,
        )
        or re.search(rf'\bwithout\s+(?:(?:an?|the)\s+)?{output_format}\b', clause)
    )


def generated_file_export_requested(user_question: str) -> bool:
    """Return whether the user asked for a supported generated file artifact."""
    return get_requested_generated_file_format(user_question) is not None


def _question_requires_derived_output(normalized_question: str) -> bool:
    return any(pattern.search(normalized_question) for pattern in PASSTHROUGH_DERIVED_OUTPUT_PATTERNS)


def _question_requests_unchanged_copy(normalized_question: str) -> bool:
    return any(pattern.search(normalized_question) for pattern in PASSTHROUGH_COPY_PATTERNS)


def _question_requests_serialization(normalized_question: str) -> bool:
    return any(pattern.search(normalized_question) for pattern in PASSTHROUGH_SERIALIZE_PATTERNS)


def _collect_row_schema(rows: Sequence[Dict[str, Any]]) -> List[str]:
    for row in rows or []:
        if isinstance(row, dict):
            return [str(field_name or '').strip() for field_name in row if str(field_name or '').strip()]
    return []


def evaluate_generated_file_passthrough_eligibility(
    user_question: str,
    rows: Optional[Sequence[Dict[str, Any]]] = None,
    public_output_schema: Optional[Sequence[str]] = None,
) -> Dict[str, Any]:
    """Return whether raw rows can satisfy a requested generated file contract."""
    normalized_question = _normalize_question_for_artifact_detection(user_question)
    normalized_rows = [row for row in list(rows or []) if isinstance(row, dict)]
    if not normalized_rows:
        return {'allowed': False, 'reason_code': 'source_result_incomplete'}

    normalized_public_schema = [
        str(field_name or '').strip()
        for field_name in list(public_output_schema or [])
        if str(field_name or '').strip()
    ]
    if normalized_public_schema:
        expected_schema = set(normalized_public_schema)
        if _collect_row_schema(normalized_rows) != normalized_public_schema:
            return {'allowed': False, 'reason_code': 'schema_not_satisfied'}
        if any(set(row.keys()) != expected_schema for row in normalized_rows):
            return {'allowed': False, 'reason_code': 'schema_not_satisfied'}

    if _question_requests_unchanged_copy(normalized_question):
        return {'allowed': True, 'reason_code': 'explicit_unchanged_copy'}
    if _question_requires_derived_output(normalized_question):
        return {'allowed': False, 'reason_code': 'derived_output_requires_transform'}
    if _question_requests_serialization(normalized_question):
        return {'allowed': True, 'reason_code': 'explicit_format_conversion'}
    return {'allowed': False, 'reason_code': 'no_explicit_passthrough_contract'}


def build_generated_file_output_guidance(
    user_question: str,
    requested_format: Optional[str] = None,
) -> str:
    """Return shared model guidance for a requested generated output format."""
    output_format = (
        str(requested_format or '').strip().lower()
        or get_requested_generated_file_format(user_question)
        or get_requested_structured_artifact_format(user_question)
    )
    if output_format == GENERATED_FILE_FORMAT_CSV:
        return build_csv_output_clarification_guidance(user_question)
    if output_format == 'json':
        return (
            'The user requested a downloadable JSON artifact. The server will validate and attach the file after '
            'generation. Return ONLY the complete valid JSON payload needed for that file. Do not wrap it in Markdown, '
            'add explanations, claim that files cannot be attached, tell the user to copy or save content manually, '
            'or mention the publication mechanism.'
        )
    if output_format == 'xml':
        return (
            'The user requested a downloadable XML artifact. The server will validate and attach the file after '
            'generation. Return ONLY one complete well-formed XML document needed for that file. Do not wrap it in '
            'Markdown, add explanations, claim that files cannot be attached, tell the user to copy or save content '
            'manually, or mention the publication mechanism.'
        )
    if output_format in {GENERATED_FILE_FORMAT_DOCX, GENERATED_FILE_FORMAT_PDF}:
        return (
            f'The user requested a downloadable {output_format.upper()} artifact. Provide a clear final '
            'response grounded in the available evidence. Structured function results from this turn may '
            'be included as labeled tables in the generated file; do not invent rows or claim an attachment '
            'exists before the file-output finalizer publishes it.'
        )
    return ''


def get_generated_file_export_content(assistant_result: Any) -> str:
    """Return the structured document-action reply when it supersedes a concise artifact reply."""
    if not isinstance(assistant_result, dict):
        return str(assistant_result or '')

    analysis_result = assistant_result.get('analysis_result')
    if isinstance(analysis_result, dict):
        analysis_reply = str(analysis_result.get('analysis_reply') or '').strip()
        if analysis_reply:
            return analysis_reply

    return str(assistant_result.get('reply') or '')


def build_generated_file_export(
    user_question: str,
    assistant_content: str,
    function_results: Optional[List[Dict[str, Any]]] = None,
) -> Optional[Dict[str, Any]]:
    """Build a generated file payload from final assistant content and function-result evidence."""
    output_format = get_requested_generated_file_format(user_question)
    if output_format not in GENERATED_FILE_FORMATS:
        return None

    assistant_text = str(assistant_content or '').strip()
    assistant_rows = extract_assistant_table_entries(assistant_text)
    function_rows = extract_authorized_function_result_rows(function_results)
    function_passthrough = evaluate_generated_file_passthrough_eligibility(
        user_question,
        rows=function_rows,
    ) if function_rows else {'allowed': False, 'reason_code': 'source_result_incomplete'}

    if output_format == GENERATED_FILE_FORMAT_CSV:
        rows = assistant_rows or (function_rows if function_passthrough.get('allowed') else [])
        if not rows:
            return None
        row_source = 'assistant response' if assistant_rows else 'structured function result'
        return _build_generated_file_payload(
            output_format=output_format,
            file_content=build_assistant_table_csv(rows),
            rows=rows,
            row_source=row_source,
            assistant_content=assistant_text,
            passthrough_reason_code=(None if assistant_rows else function_passthrough.get('reason_code')),
        )

    if not assistant_text and not assistant_rows and not function_rows:
        return None
    rows = function_rows if function_rows and function_passthrough.get('allowed') else assistant_rows
    if not assistant_text and not rows:
        return None
    row_source = 'structured function result' if rows and rows is function_rows else 'assistant response'
    title = _build_generated_file_title(output_format)
    if output_format == GENERATED_FILE_FORMAT_DOCX:
        file_content = _render_docx_file_export(title, assistant_text, rows, row_source)
    else:
        file_content = _render_pdf_file_export(title, assistant_text, rows, row_source)
    return _build_generated_file_payload(
        output_format=output_format,
        file_content=file_content,
        rows=rows,
        row_source=row_source,
        assistant_content=assistant_text,
        title=title,
        passthrough_reason_code=(function_passthrough.get('reason_code') if row_source == 'structured function result' else None),
    )


def extract_authorized_function_result_rows(function_results: Optional[List[Dict[str, Any]]]) -> List[Dict[str, Any]]:
    """Return structured rows from successful current-turn non-tabular function results."""
    function_row_groups: List[Tuple[str, List[Dict[str, Any]]]] = []
    for function_result in function_results or []:
        if not isinstance(function_result, dict):
            continue
        if function_result.get('success') is False or _is_tabular_function_result(function_result):
            continue

        structured_rows = _extract_function_result_rows(
            _parse_function_result_payload(function_result.get('function_result')),
        )
        if not structured_rows:
            continue
        function_row_groups.append((
            _get_function_result_label(function_result),
            structured_rows,
        ))

    if not function_row_groups:
        return []
    if len(function_row_groups) == 1:
        return function_row_groups[0][1]

    source_column = _get_function_result_source_column(function_row_groups)
    combined_rows = []
    for function_label, rows in function_row_groups:
        for row in rows:
            normalized_row = dict(row)
            normalized_row[source_column] = function_label
            combined_rows.append(normalized_row)
    return combined_rows


def has_generated_file_output(existing_outputs: Optional[List[Dict[str, Any]]], output_format: str) -> bool:
    """Return whether an existing generated artifact already covers an output format."""
    normalized_output_format = str(output_format or '').strip().lower()
    if not normalized_output_format:
        return False
    for output in existing_outputs or []:
        if not isinstance(output, dict):
            continue
        existing_output_format = str(output.get('output_format') or '').strip().lower()
        existing_file_name = str(output.get('file_name') or '').strip().lower()
        if existing_output_format == normalized_output_format or existing_file_name.endswith(f'.{normalized_output_format}'):
            return True
    return False


def build_generated_file_artifact_metadata(
    export_payload: Dict[str, Any],
    upload_result: Dict[str, Any],
    conversation_id: str,
) -> Optional[Dict[str, Any]]:
    """Build public artifact metadata after an authorized generated-file upload."""
    uploaded_message = upload_result.get('message') if isinstance(upload_result, dict) else {}
    uploaded_message = uploaded_message if isinstance(uploaded_message, dict) else {}
    artifact_message_id = str(uploaded_message.get('id') or '').strip()
    if not artifact_message_id:
        return None

    generated_file_name = str(export_payload.get('file_name') or '').strip()
    normalized_output_format = str(export_payload.get('output_format') or '').strip().lower()
    artifact_metadata = {
        'capability': str(export_payload.get('capability') or 'file_export').strip().lower() or 'file_export',
        'artifact_message_id': artifact_message_id,
        'conversation_id': str(conversation_id or '').strip(),
        'storage_scope': 'chat',
        'file_name': uploaded_message.get('file_name') or generated_file_name,
        'output_format': normalized_output_format,
        'summary': str(export_payload.get('summary') or '').strip(),
        'suppress_assistant_text': normalized_output_format in ASSISTANT_TEXT_SUPPRESSING_FORMATS,
    }
    row_count = export_payload.get('row_count')
    if isinstance(row_count, int) and row_count > 0:
        artifact_metadata['row_count'] = row_count
    preview_rows = export_payload.get('preview_rows')
    if isinstance(preview_rows, list) and preview_rows:
        artifact_metadata['preview_rows'] = preview_rows
        if isinstance(preview_rows[0], dict):
            artifact_metadata['preview_columns'] = list(preview_rows[0])[:50]
    preview_lines = export_payload.get('preview_lines')
    if isinstance(preview_lines, list) and preview_lines:
        artifact_metadata['preview_lines'] = preview_lines
    row_source = str(export_payload.get('row_source') or '').strip()
    if row_source:
        artifact_metadata['row_source'] = row_source
    return artifact_metadata


def _build_generated_file_payload(
    output_format: str,
    file_content: Any,
    rows: Sequence[Dict[str, Any]],
    row_source: str,
    assistant_content: str,
    title: str = '',
    passthrough_reason_code: Optional[str] = None,
) -> Dict[str, Any]:
    normalized_output_format = str(output_format or '').strip().lower()
    row_count = len(rows or [])
    normalized_title = str(title or _build_generated_file_title(normalized_output_format)).strip()
    payload = {
        'capability': 'file_export',
        'file_name': _build_generated_file_name(normalized_output_format),
        'file_content': file_content,
        'output_format': normalized_output_format,
        'row_count': row_count,
        'preview_rows': list(rows or [])[:GENERATED_FILE_PREVIEW_ROWS],
        'preview_lines': _build_preview_lines(assistant_content),
        'row_source': row_source,
        '_structured_rows': list(rows or []),
        'summary': _build_generated_file_summary(
            normalized_output_format,
            row_count,
            row_source,
            normalized_title,
        ),
    }
    if passthrough_reason_code:
        payload['passthrough_reason_code'] = str(passthrough_reason_code or '').strip()[:80]
    return payload


def _build_generated_file_name(output_format: str) -> str:
    timestamp_suffix = datetime.utcnow().strftime('%Y%m%d_%H%M%S')
    return f'generated_output_{timestamp_suffix}.{output_format}'


def _build_generated_file_title(output_format: str) -> str:
    return f'Generated {str(output_format or "file").upper()} export'


def _build_generated_file_summary(
    output_format: str,
    row_count: int,
    row_source: str,
    title: str,
) -> str:
    row_detail = f' with {row_count} structured row(s)' if row_count else ''
    return f'Prepared {title}{row_detail} from the {row_source}.'


def _build_preview_lines(assistant_content: str) -> List[str]:
    normalized_lines = [
        line.strip()
        for line in str(assistant_content or '').splitlines()
        if line.strip()
    ]
    return normalized_lines[:3]


def _parse_function_result_payload(value: Any) -> Any:
    if not isinstance(value, str):
        return value

    normalized_value = value.strip()
    if not normalized_value or normalized_value[0] not in '[{':
        return value
    try:
        return json.loads(normalized_value)
    except (TypeError, ValueError, json.JSONDecodeError):
        return value


def _extract_function_result_rows(
    payload: Any,
    depth: int = 0,
    is_data_row: bool = False,
) -> List[Dict[str, Any]]:
    if depth > 4:
        return []
    if isinstance(payload, list):
        rows = []
        for item in payload:
            if isinstance(item, dict):
                normalized_row = _normalize_function_result_row(item, is_data_row=True)
                if normalized_row:
                    rows.append(normalized_row)
            elif item not in (None, ''):
                rows.append({'value': _sanitize_function_result_value(item)})
        return rows
    if not isinstance(payload, dict):
        return []

    normalized_keys = {
        _normalize_function_result_key(key): key
        for key in payload
    }
    for row_key in FUNCTION_RESULT_ROW_KEYS:
        matching_key = normalized_keys.get(_normalize_function_result_key(row_key))
        if matching_key is None:
            continue
        rows = _extract_function_result_rows(
            payload.get(matching_key),
            depth + 1,
            is_data_row=True,
        )
        if rows:
            return rows

    normalized_row = _normalize_function_result_row(payload, is_data_row=is_data_row)
    return [normalized_row] if normalized_row else []


def _normalize_function_result_row(row: Dict[str, Any], is_data_row: bool) -> Dict[str, Any]:
    normalized_row = {}
    for raw_key, raw_value in row.items():
        key = str(raw_key or '').strip()
        normalized_key = _normalize_function_result_key(key)
        if not key or _is_sensitive_function_result_key(key):
            continue
        if not is_data_row and normalized_key in FUNCTION_RESULT_CONTROL_KEYS:
            continue

        value = _sanitize_function_result_value(raw_value)
        if value in (None, '', [], {}) or value == '***REDACTED***':
            continue
        normalized_row[key] = value
    return normalized_row


def _sanitize_function_result_value(value: Any, depth: int = 0) -> Any:
    if depth > 4:
        return '[truncated]'
    if isinstance(value, dict):
        return {
            str(key): _sanitize_function_result_value(item, depth + 1)
            for key, item in value.items()
            if not _is_sensitive_function_result_key(key)
        }
    if isinstance(value, (list, tuple, set)):
        return [
            _sanitize_function_result_value(item, depth + 1)
            for item in value
        ]
    return value


def _normalize_function_result_key(key: Any) -> str:
    return re.sub(r'[^a-z0-9]', '', str(key or '').casefold())


def _is_sensitive_function_result_key(key: Any) -> bool:
    normalized_key = _normalize_function_result_key(key)
    if not normalized_key:
        return False
    return any(fragment in normalized_key for fragment in FUNCTION_RESULT_SENSITIVE_KEY_FRAGMENTS)


def _is_tabular_function_result(function_result: Dict[str, Any]) -> bool:
    plugin_name = str(function_result.get('plugin_name') or '').strip().casefold()
    return plugin_name in TABULAR_FUNCTION_RESULT_PLUGIN_NAMES


def _get_function_result_label(function_result: Dict[str, Any]) -> str:
    return (
        str(function_result.get('function_name') or '').strip()
        or str(function_result.get('plugin_name') or '').strip()
        or 'function result'
    )


def _get_function_result_source_column(
    function_row_groups: Sequence[Tuple[str, Sequence[Dict[str, Any]]]],
) -> str:
    existing_columns = {
        str(column_name).casefold()
        for _, rows in function_row_groups
        for row in rows
        for column_name in row
    }
    source_column = 'Source action'
    suffix = 2
    while source_column.casefold() in existing_columns:
        source_column = f'Source action {suffix}'
        suffix += 1
    return source_column


def _render_docx_file_export(
    title: str,
    assistant_content: str,
    rows: Sequence[Dict[str, Any]],
    row_source: str,
) -> bytes:
    from docx import Document as DocxDocument

    document = DocxDocument()
    document.add_heading(title, level=1)
    _append_docx_text(document, assistant_content)
    if rows:
        document.add_heading(_build_structured_rows_heading(row_source), level=2)
        _append_docx_table(document, rows)

    output_buffer = io.BytesIO()
    document.save(output_buffer)
    return output_buffer.getvalue()


def _append_docx_text(document: Any, assistant_content: str) -> None:
    normalized_content = str(assistant_content or '').strip()
    if not normalized_content:
        return
    for paragraph_text in re.split(r'\n\s*\n', normalized_content):
        cleaned_paragraph = paragraph_text.strip()
        if cleaned_paragraph:
            document.add_paragraph(cleaned_paragraph)


def _append_docx_table(document: Any, rows: Sequence[Dict[str, Any]]) -> None:
    columns = _collect_structured_row_columns(rows)
    if not columns:
        return
    table = document.add_table(rows=1, cols=len(columns))
    table.style = 'Table Grid'
    for index, column_name in enumerate(columns):
        table.rows[0].cells[index].text = str(column_name)
    for row in rows:
        cells = table.add_row().cells
        for index, column_name in enumerate(columns):
            cells[index].text = _format_structured_cell(row.get(column_name))


def _render_pdf_file_export(
    title: str,
    assistant_content: str,
    rows: Sequence[Dict[str, Any]],
    row_source: str,
) -> bytes:
    import fitz

    html_parts = [f'<h1>{html.escape(title)}</h1>']
    normalized_content = str(assistant_content or '').strip()
    if normalized_content:
        html_parts.append('<h2>Response</h2>')
        for paragraph_text in re.split(r'\n\s*\n', normalized_content):
            cleaned_paragraph = paragraph_text.strip()
            if cleaned_paragraph:
                html_parts.append(f'<p>{html.escape(cleaned_paragraph).replace(chr(10), "<br/>")}</p>')
    if rows:
        html_parts.append(f'<h2>{html.escape(_build_structured_rows_heading(row_source))}</h2>')
        html_parts.append(_build_structured_rows_html(rows))

    media_box = fitz.paper_rect('letter')
    content_box = media_box + (36, 36, -36, -36)
    story = fitz.Story(html='\n'.join(html_parts), user_css=PDF_EXPORT_CSS)
    temporary_path = None
    try:
        with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as temporary_file:
            temporary_path = temporary_file.name

        writer = fitz.DocumentWriter(temporary_path)
        has_more = True
        while has_more:
            device = writer.begin_page(media_box)
            has_more, _ = story.place(content_box)
            story.draw(device)
            writer.end_page()
        writer.close()
        with open(temporary_path, 'rb') as generated_file:
            return generated_file.read()
    finally:
        if temporary_path:
            try:
                os.unlink(temporary_path)
            except OSError:
                pass


def _build_structured_rows_html(rows: Sequence[Dict[str, Any]]) -> str:
    columns = _collect_structured_row_columns(rows)
    if not columns:
        return '<p><i>No structured rows were available.</i></p>'
    table_parts = ['<table><tr>']
    table_parts.extend(f'<th>{html.escape(str(column_name))}</th>' for column_name in columns)
    table_parts.append('</tr>')
    for row in rows:
        table_parts.append('<tr>')
        for column_name in columns:
            table_parts.append(f'<td>{html.escape(_format_structured_cell(row.get(column_name))).replace(chr(10), "<br/>")}</td>')
        table_parts.append('</tr>')
    table_parts.append('</table>')
    return ''.join(table_parts)


def _collect_structured_row_columns(rows: Sequence[Dict[str, Any]]) -> List[str]:
    columns = []
    seen_columns = set()
    for row in rows or []:
        if not isinstance(row, dict):
            continue
        for raw_column_name in row:
            column_name = str(raw_column_name or '').strip()
            if not column_name or column_name.casefold() in seen_columns:
                continue
            seen_columns.add(column_name.casefold())
            columns.append(column_name)
    return columns


def _format_structured_cell(value: Any) -> str:
    if value is None:
        return ''
    if isinstance(value, (dict, list, tuple, set)):
        return json.dumps(value, default=str, ensure_ascii=False)
    return str(value)


def _build_structured_rows_heading(row_source: str) -> str:
    if row_source == 'structured function result':
        return 'Structured function results'
    return 'Structured response rows'


def normalize_generated_output_format(output_format, default='json'):
    """Normalize generated artifact output formats supported by the export framework."""
    normalized_format = str(output_format or '').strip().lower().lstrip('.')
    if normalized_format in SUPPORTED_GENERATED_EXPORT_FORMATS:
        return normalized_format

    normalized_default = str(default or 'json').strip().lower().lstrip('.')
    if normalized_default in SUPPORTED_GENERATED_EXPORT_FORMATS:
        return normalized_default
    return 'json'


def strip_markdown_code_fence(text):
    """Remove a single surrounding Markdown code fence while preserving content."""
    normalized_text = str(text or '').strip()
    if not normalized_text.startswith('```'):
        return normalized_text

    header_end_index = normalized_text.find('\n')
    if header_end_index <= 0:
        return normalized_text

    header_suffix = normalized_text[3:header_end_index].strip()
    if header_suffix and not all(character.isalnum() or character in {'_', '-'} for character in header_suffix):
        return normalized_text

    closing_index = normalized_text.rfind('```')
    if closing_index <= header_end_index:
        return normalized_text

    trailing_text = normalized_text[closing_index + 3:].strip()
    if trailing_text:
        return normalized_text

    return normalized_text[header_end_index + 1:closing_index].strip()


def _iter_xml_candidates(text) -> Iterable[str]:
    normalized_text = strip_markdown_code_fence(text)
    if not normalized_text:
        return

    yield normalized_text

    first_xml_index = normalized_text.find('<?xml')
    if first_xml_index > 0:
        yield normalized_text[first_xml_index:].strip()

    first_tag_index = normalized_text.find('<')
    if first_tag_index > 0:
        yield normalized_text[first_tag_index:].strip()

    for root_match in XML_ROOT_PATTERN.finditer(normalized_text):
        root_tag = root_match.group('tag')
        root_start = root_match.start()
        root_open = root_match.group(0)
        if root_open.rstrip().endswith('/>'):
            yield normalized_text[root_start:root_match.end()].strip()
            continue

        closing_tag = f'</{root_tag}>'
        root_end = normalized_text.rfind(closing_tag)
        if root_end <= root_start:
            continue

        yield normalized_text[root_start:root_end + len(closing_tag)].strip()


def normalize_xml_artifact_payload(text):
    """Return a complete XML document extracted from model output, or an empty string."""
    seen_candidates = set()
    for candidate in _iter_xml_candidates(text):
        if candidate in seen_candidates:
            continue
        seen_candidates.add(candidate)
        try:
            DefusedElementTree.fromstring(candidate.encode('utf-8'))
        except (DefusedXmlException, ElementTree.ParseError):
            continue
        return candidate
    return ''


def normalize_json_artifact_payload(text):
    """Return parsed JSON extracted from model output, or None when no JSON is present."""
    normalized_text = strip_markdown_code_fence(text)
    if not normalized_text:
        return None

    decoder = json.JSONDecoder()
    try:
        parsed_value, _ = decoder.raw_decode(normalized_text)
        return parsed_value
    except (TypeError, ValueError, json.JSONDecodeError):
        pass

    for start_index, character in enumerate(normalized_text):
        if character not in '[{':
            continue
        try:
            parsed_value, _ = decoder.raw_decode(normalized_text[start_index:])
            return parsed_value
        except (TypeError, ValueError, json.JSONDecodeError):
            continue

    return None


def _sanitize_xml_tag_name(value, fallback_value):
    normalized_name = re.sub(r'[^A-Za-z0-9_.-]+', '_', str(value or '').strip())
    normalized_name = normalized_name.strip('._-')
    if not normalized_name:
        normalized_name = fallback_value
    if not re.match(r'^[A-Za-z_]', normalized_name):
        normalized_name = f'{fallback_value}_{normalized_name}'
    return normalized_name


def _append_xml_value(parent, value, item_name):
    if isinstance(value, dict):
        for key, child_value in value.items():
            child = ElementTree.SubElement(
                parent,
                _sanitize_xml_tag_name(key, 'Field'),
            )
            _append_xml_value(child, child_value, item_name)
        return

    if isinstance(value, (list, tuple)):
        for item in value:
            child = ElementTree.SubElement(
                parent,
                _sanitize_xml_tag_name(item_name, 'Item'),
            )
            _append_xml_value(child, item, item_name)
        return

    if value is None:
        parent.text = ''
        return

    if isinstance(value, bool):
        parent.text = 'true' if value else 'false'
        return

    parent.text = str(value)


def build_xml_from_value(value: Any, root_name='GeneratedOutput', item_name='Item'):
    """Serialize a Python value into a deterministic XML document."""
    root = ElementTree.Element(_sanitize_xml_tag_name(root_name, 'GeneratedOutput'))
    _append_xml_value(root, value, item_name)
    ElementTree.indent(root, space='  ')
    xml_body = ElementTree.tostring(root, encoding='unicode', short_empty_elements=True)
    return f'{XML_DECLARATION}\n{xml_body}'


def serialize_generated_xml(value: Any, root_name='GeneratedOutput', item_name='Item'):
    """Serialize generated content to XML, preserving valid XML model output when present."""
    if isinstance(value, str):
        xml_payload = normalize_xml_artifact_payload(value)
        if xml_payload:
            return xml_payload

    return build_xml_from_value(value, root_name=root_name, item_name=item_name)


def serialize_generated_json(value: Any, *, indent=2):
    """Serialize generated content to JSON using the export framework defaults."""
    return json.dumps(value, indent=indent, ensure_ascii=False, default=str)

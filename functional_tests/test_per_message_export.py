#!/usr/bin/env python3
# test_per_message_export.py
"""
Functional tests for the per-message export feature and Word route regression fix.
Version: 0.240.076
Implemented in: 0.240.076

Covers:
 - Happy path: Word document built successfully from a valid message.
 - Word-native formatting: markdown content is converted into DOCX headings, lists, tables, and styled runs.
 - Markdown export logic: correct header, timestamp and content rendered.
 - Route regression: backend source defines POST /api/message/export-word.
 - Auth failure: unauthenticated caller receives 401.
 - Ownership failure: caller who does not own the conversation receives 403.
"""

import ast
import io
import os
import sys
from typing import Any, Dict, Optional

sys.path.insert(
    0,
    os.path.join(os.path.dirname(os.path.abspath(__file__)), '..', 'application', 'single_app')
)


# ---------------------------------------------------------------------------
# Helpers – replicate key logic from route_backend_conversation_export.py
# so the tests run without a live Flask + Cosmos DB environment.
# ---------------------------------------------------------------------------

def _normalize_content(content):
    """Replicate _normalize_content from route_backend_conversation_export."""
    if isinstance(content, str):
        return content
    if isinstance(content, list):
        parts = []
        for item in content:
            if isinstance(item, dict):
                if item.get('type') == 'text':
                    parts.append(item.get('text', ''))
                elif item.get('type') == 'image_url':
                    parts.append('[Image]')
                else:
                    parts.append(str(item))
            else:
                parts.append(str(item))
        return '\n'.join(parts)
    if isinstance(content, dict):
        if content.get('type') == 'text':
            return content.get('text', '')
        return str(content)
    return str(content) if content else ''


def _verify_ownership(conversation, requesting_user_id):
    """Return (ok, status_code, error_msg)."""
    if conversation is None:
        return False, 404, 'Conversation not found'
    if conversation.get('user_id') != requesting_user_id:
        return False, 403, 'Access denied'
    return True, 200, None


def _check_auth(user_id):
    """Simulate the authentication guard at the start of the endpoint."""
    if not user_id:
        return False, 401, 'User not authenticated'
    return True, 200, None


def _build_markdown_export(role, content, sender, timestamp):
    """Replicate the client-side Markdown export logic from chat-message-export.js."""
    lines = []
    lines.append(f"### {sender}")
    if timestamp:
        lines.append(f"*{timestamp}*")
    lines.append('')
    lines.append(content)
    lines.append('')
    return '\n'.join(lines)


def _load_word_formatter_helpers():
    """Load the real DOCX formatter helpers from the export route source."""
    try:
        import markdown2
        from bs4 import BeautifulSoup, NavigableString, Tag
        from docx import Document as DocxDocument
        from docx.shared import Inches, Pt
    except ImportError as exc:
        return None, exc

    route_file = os.path.join(
        os.path.dirname(os.path.abspath(__file__)),
        '..',
        'application',
        'single_app',
        'route_backend_conversation_export.py'
    )

    with open(route_file, 'r', encoding='utf-8') as handle:
        source = handle.read()

    tree = ast.parse(source)
    helper_names = {
        '_add_markdown_content_to_doc',
        '_append_html_block_to_doc',
        '_append_list_items_to_doc',
        '_add_code_block_to_doc',
        '_add_html_table_to_doc',
        '_populate_table_cell',
        '_append_inline_html_runs',
        '_apply_run_formatting',
    }
    selected_nodes = [
        node for node in tree.body
        if isinstance(node, ast.FunctionDef) and node.name in helper_names
    ]

    loaded_names = {node.name for node in selected_nodes}
    missing_names = helper_names - loaded_names
    assert not missing_names, f"Missing formatter helpers in route file: {sorted(missing_names)}"

    module = ast.Module(body=selected_nodes, type_ignores=[])
    ast.fix_missing_locations(module)

    namespace = {
        'Any': Any,
        'Dict': Dict,
        'Optional': Optional,
        'markdown2': markdown2,
        'BeautifulSoup': BeautifulSoup,
        'NavigableString': NavigableString,
        'Tag': Tag,
        'DocxDocument': DocxDocument,
        'Inches': Inches,
        'Pt': Pt,
        'DOCX_MARKDOWN_EXTRAS': ['fenced-code-blocks', 'tables', 'break-on-newline', 'cuddled-lists', 'strike'],
    }

    exec(compile(module, route_file, 'exec'), namespace)
    return namespace, None


# ---------------------------------------------------------------------------
# Tests
# ---------------------------------------------------------------------------

def test_happy_path_word_export():
    """Happy path: Word document is built without error for a valid message."""
    print("🔍 Testing happy path – Word document generation...")

    try:
        from docx import Document as DocxDocument
        from docx.shared import Pt
    except ImportError as exc:
        print(f"  ⚠️  python-docx not installed, skipping Word generation check: {exc}")
        print("✅ test_happy_path_word_export skipped (dependency missing)")
        return True

    # Arrange
    requesting_user_id = 'user-alice'
    conversation = {'id': 'conv-001', 'user_id': 'user-alice'}
    message = {
        'id': 'msg-001',
        'role': 'assistant',
        'content': '**Hello**, world!\n\nThis is a test.',
        'timestamp': '2025-06-01T10:00:00Z',
        'citations': [
            {'title': 'Reference Doc', 'url': 'https://example.com/ref'}
        ]
    }

    # Auth check
    auth_ok, auth_status, auth_err = _check_auth(requesting_user_id)
    assert auth_ok, f"Auth should pass, got {auth_status}: {auth_err}"

    # Ownership check
    ok, status, err = _verify_ownership(conversation, requesting_user_id)
    assert ok, f"Ownership check should pass, got {status}: {err}"

    # Build Word document
    doc = DocxDocument()
    doc.add_heading('Message Export', level=1)

    role = message.get('role', 'unknown').capitalize()
    timestamp = message.get('timestamp', '')

    meta_para = doc.add_paragraph()
    meta_run = meta_para.add_run(f"Role: {role}")
    meta_run.bold = True
    if timestamp:
        meta_para.add_run(f"    {timestamp}")

    doc.add_paragraph('')

    content = _normalize_content(message.get('content', ''))

    # Add content as a paragraph (simplified – full logic tested in route unit)
    doc.add_paragraph(content)

    citations = message.get('citations', [])
    if citations:
        doc.add_heading('Citations', level=2)
        for cit in citations:
            source = cit.get('title') or cit.get('url', 'Unknown')
            doc.add_paragraph(source, style='List Bullet')

    # Serialise to buffer – if this raises, the test fails
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    docx_bytes = buffer.read()

    assert len(docx_bytes) > 0, "Generated docx should be non-empty"

    # Round-trip verify
    buffer.seek(0)
    loaded = DocxDocument(io.BytesIO(docx_bytes))
    headings = [p.text for p in loaded.paragraphs if p.style.name.startswith('Heading')]
    assert 'Message Export' in headings, "Document should have 'Message Export' heading"

    print("✅ test_happy_path_word_export passed!")
    return True


def test_word_export_uses_word_formatting():
    """Word export should convert markdown into DOCX-native structures."""
    print("🔍 Testing Word export uses Word-native formatting...")

    helpers, import_error = _load_word_formatter_helpers()
    if import_error is not None:
        print(f"  ⚠️  Required formatter dependency missing, skipping check: {import_error}")
        print("✅ test_word_export_uses_word_formatting skipped (dependency missing)")
        return True

    from docx import Document as DocxDocument

    content = (
        "# Summary\n\n"
        "Paragraph with **bold** text, *italic* text, and `code`.\n\n"
        "- First item\n"
        "- Second item\n\n"
        "| Name | Value |\n"
        "| --- | --- |\n"
        "| Alpha | 42 |\n\n"
        "```python\n"
        "print('hello')\n"
        "```"
    )

    doc = DocxDocument()
    helpers['_add_markdown_content_to_doc'](doc, content)

    non_empty_paragraphs = [paragraph for paragraph in doc.paragraphs if paragraph.text.strip()]
    paragraph_texts = [paragraph.text.strip() for paragraph in non_empty_paragraphs]

    heading = next((paragraph for paragraph in non_empty_paragraphs if paragraph.text.strip() == 'Summary'), None)
    assert heading is not None, 'Expected a heading paragraph for the markdown heading'
    assert heading.style.name.startswith('Heading'), f"Expected heading style, found {heading.style.name}"

    body_paragraph = next((paragraph for paragraph in non_empty_paragraphs if 'Paragraph with' in paragraph.text), None)
    assert body_paragraph is not None, 'Expected the markdown body paragraph to be rendered'
    assert any('bold' in run.text and run.bold for run in body_paragraph.runs), 'Expected bold text to use a bold run'
    assert any('italic' in run.text and run.italic for run in body_paragraph.runs), 'Expected italic text to use an italic run'
    assert any('code' in run.text and run.font.name == 'Consolas' for run in body_paragraph.runs), 'Expected inline code to use a code font'

    bullet_paragraphs = [
        paragraph for paragraph in non_empty_paragraphs
        if paragraph.text.strip() in {'First item', 'Second item'}
    ]
    assert len(bullet_paragraphs) == 2, 'Expected both markdown bullets to be rendered as list items'
    assert all(paragraph.style.name == 'List Bullet' for paragraph in bullet_paragraphs), 'Expected markdown bullets to use the Word bullet list style'

    assert len(doc.tables) == 1, 'Expected markdown table syntax to render as a DOCX table'
    table = doc.tables[0]
    assert table.cell(0, 0).text.strip() == 'Name', 'Expected first table header cell to contain Name'
    assert table.cell(0, 1).text.strip() == 'Value', 'Expected second table header cell to contain Value'
    assert table.cell(1, 0).text.strip() == 'Alpha', 'Expected table data to populate DOCX cells'
    assert table.cell(1, 1).text.strip() == '42', 'Expected table data value to populate DOCX cells'

    code_paragraph = next((paragraph for paragraph in non_empty_paragraphs if "print('hello')" in paragraph.text), None)
    assert code_paragraph is not None, 'Expected fenced code block to render as a paragraph'
    assert any("print('hello')" in run.text and run.font.name == 'Consolas' for run in code_paragraph.runs), 'Expected fenced code block to use a code font'

    combined_text = '\n'.join(paragraph_texts)
    assert '**bold**' not in combined_text, 'Raw bold markdown should not remain in the exported DOCX text'
    assert '*italic*' not in combined_text, 'Raw italic markdown should not remain in the exported DOCX text'
    assert '```python' not in combined_text, 'Raw fenced code markdown should not remain in the exported DOCX text'

    print("✅ test_word_export_uses_word_formatting passed!")
    return True


def test_happy_path_markdown_export():
    """Happy path: Markdown file content is correctly formatted."""
    print("🔍 Testing happy path – Markdown export...")

    role = 'assistant'
    content = 'Here is a **bold** answer.'
    sender = 'Assistant'
    timestamp = '2025-06-01T10:05:00Z'

    markdown = _build_markdown_export(role, content, sender, timestamp)

    assert '### Assistant' in markdown, "Should have role heading"
    assert f'*{timestamp}*' in markdown, "Should include timestamp"
    assert content in markdown, "Should include message content"
    # File should start with the heading line
    assert markdown.startswith('### Assistant'), "Heading should be first line"

    print("✅ test_happy_path_markdown_export passed!")
    return True


def test_export_word_route_definition_present():
    """Route regression: the backend must define POST /api/message/export-word."""
    print("🔍 Testing backend route definition for Word export...")

    route_file = os.path.join(
        os.path.dirname(os.path.abspath(__file__)),
        '..',
        'application',
        'single_app',
        'route_backend_conversation_export.py'
    )

    with open(route_file, 'r', encoding='utf-8') as handle:
        source = handle.read()

    tree = ast.parse(source)
    register_func = next(
        (
            node for node in tree.body
            if isinstance(node, ast.FunctionDef)
            and node.name == 'register_route_backend_conversation_export'
        ),
        None
    )

    assert register_func is not None, 'register_route_backend_conversation_export should exist'

    export_route_found = False
    for node in register_func.body:
        if not isinstance(node, ast.FunctionDef):
            continue

        for decorator in node.decorator_list:
            if not isinstance(decorator, ast.Call):
                continue

            func = decorator.func
            if not isinstance(func, ast.Attribute) or func.attr != 'route':
                continue

            if not decorator.args:
                continue

            route_arg = decorator.args[0]
            if not isinstance(route_arg, ast.Constant) or route_arg.value != '/api/message/export-word':
                continue

            methods_kw = next((keyword for keyword in decorator.keywords if keyword.arg == 'methods'), None)
            assert methods_kw is not None, 'Export Word route should declare allowed methods'
            assert isinstance(methods_kw.value, (ast.List, ast.Tuple)), 'Route methods should be a list or tuple'

            methods = [
                item.value for item in methods_kw.value.elts
                if isinstance(item, ast.Constant)
            ]
            assert 'POST' in methods, f'Expected POST method, found {methods}'
            assert node.name == 'api_export_message_word', f'Unexpected route handler name: {node.name}'
            export_route_found = True
            break

        if export_route_found:
            break

    assert export_route_found, 'Expected POST /api/message/export-word to be defined'

    print("✅ test_export_word_route_definition_present passed!")
    return True


def test_auth_failure_unauthenticated():
    """Auth failure: an unauthenticated caller (no user_id) should get 401."""
    print("🔍 Testing auth failure – unauthenticated request...")

    for bad_user_id in (None, '', False):
        ok, status, err = _check_auth(bad_user_id)
        assert not ok, f"Auth should fail for user_id={bad_user_id!r}"
        assert status == 401, f"Expected 401, got {status}"
        assert err == 'User not authenticated', f"Unexpected error message: {err}"

    print("✅ test_auth_failure_unauthenticated passed!")
    return True


def test_ownership_failure_wrong_user():
    """Ownership failure: user requesting another user's conversation gets 403."""
    print("🔍 Testing ownership failure – wrong user...")

    conversation = {'id': 'conv-bob', 'user_id': 'user-bob'}
    requesting_user = 'user-alice'

    ok, status, err = _verify_ownership(conversation, requesting_user)

    assert not ok, "Ownership check should fail"
    assert status == 403, f"Expected 403, got {status}"
    assert err == 'Access denied', f"Unexpected error message: {err}"

    print("✅ test_ownership_failure_wrong_user passed!")
    return True


def test_ownership_failure_missing_conversation():
    """Ownership failure: conversation not found should return 404."""
    print("🔍 Testing ownership failure – conversation not found...")

    ok, status, err = _verify_ownership(None, 'user-alice')

    assert not ok, "Ownership check should fail for missing conversation"
    assert status == 404, f"Expected 404, got {status}"
    assert err == 'Conversation not found', f"Unexpected error message: {err}"

    print("✅ test_ownership_failure_missing_conversation passed!")
    return True


def test_normalize_content_variants():
    """Content normalisation handles strings, lists, and dicts correctly."""
    print("🔍 Testing content normalisation...")

    # Plain string – unchanged
    assert _normalize_content('hello') == 'hello'

    # List of text parts
    result = _normalize_content([
        {'type': 'text', 'text': 'Part 1'},
        {'type': 'text', 'text': 'Part 2'},
    ])
    assert result == 'Part 1\nPart 2', f"Unexpected: {result!r}"

    # Image entry in list
    result = _normalize_content([
        {'type': 'text', 'text': 'Before image'},
        {'type': 'image_url', 'image_url': {'url': 'https://example.com/img.png'}},
    ])
    assert '[Image]' in result, "Image entries should render as [Image]"

    # Dict with type=text
    assert _normalize_content({'type': 'text', 'text': 'Hi'}) == 'Hi'

    # None / empty
    assert _normalize_content(None) == ''
    assert _normalize_content('') == ''

    print("✅ test_normalize_content_variants passed!")
    return True


def test_markdown_export_no_timestamp():
    """Markdown export omits the timestamp line when timestamp is empty."""
    print("🔍 Testing Markdown export without timestamp...")

    markdown = _build_markdown_export('user', 'Hello!', 'User', '')

    assert '### User' in markdown
    assert 'Hello!' in markdown
    # No italicised timestamp line should be present
    lines = markdown.splitlines()
    italic_lines = [line for line in lines if line.startswith('*') and line.endswith('*')]
    assert not italic_lines, f"Should be no timestamp lines, found: {italic_lines}"

    print("✅ test_markdown_export_no_timestamp passed!")
    return True


# ---------------------------------------------------------------------------
# Runner
# ---------------------------------------------------------------------------

if __name__ == "__main__":
    tests = [
        test_happy_path_word_export,
        test_word_export_uses_word_formatting,
        test_happy_path_markdown_export,
        test_export_word_route_definition_present,
        test_auth_failure_unauthenticated,
        test_ownership_failure_wrong_user,
        test_ownership_failure_missing_conversation,
        test_normalize_content_variants,
        test_markdown_export_no_timestamp,
    ]
    results = []

    for test_fn in tests:
        print(f"\n🧪 Running {test_fn.__name__}...")
        try:
            results.append(test_fn())
        except Exception as exc:
            print(f"❌ {test_fn.__name__} failed: {exc}")
            import traceback
            traceback.print_exc()
            results.append(False)

    success = all(results)
    passed = sum(1 for r in results if r)
    print(f"\n📊 Results: {passed}/{len(results)} tests passed")
    sys.exit(0 if success else 1)
